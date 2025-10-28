from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm
import sys




# Création du document
doc = Document()

#************************************************************
# Fonctions pour créer le document
#************************************************************

def para(phrase, nom_para):
    nom_para = doc.add_paragraph(phrase)

def run(phrase, nom_para, bold_value = False):
    nom_para.add_run(phrase).bold = bold_value

def ajouter_paragraphe(doc, texte_gras1, texte_normal1, texte_gras2, texte_normal3=None, texte_gras3=None, texte_normal4=None, texte_gras4=None, texte_normal5=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
   
    run1 = para.add_run(texte_gras1)
    run1.font.size = Pt(12)
    run1.bold = True

    run2 = para.add_run(texte_normal1)
    run2.font.size = Pt(12)
    run2.bold = False

    run3 = para.add_run(texte_gras2)
    run3.font.size = Pt(12)
    run3.bold = True
   
    if texte_normal3:
        run4 = para.add_run(texte_normal3)
        run4.font.size = Pt(12)
        run4.bold = False
   
    if texte_gras3:
        run4 = para.add_run(texte_gras3)
        run4.font.size = Pt(12)
        run4.bold = True

    if texte_normal4:
        run4 = para.add_run(texte_normal4)
        run4.font.size = Pt(12)
        run4.bold = False

    if texte_gras4:
        run4 = para.add_run(texte_gras4)
        run4.font.size = Pt(12)
        run4.bold = True
   
    if texte_normal5:
        run4 = para.add_run(texte_normal5)
        run4.font.size = Pt(12)
        run4.bold = False

def ajouter_paragraphe_italique(doc, texte_gras1, texte_normal1, texte_gras2=None, texte_normal3=None, texte_gras3=None, texte_normal4=None, texte_gras4=None, texte_normal5=None, texte_gras5=None, texte_normal6=None, texte_gras6=None, texte_normal7=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
   
    run1 = para.add_run(texte_gras1)
    run1.font.size = Pt(12)
    run1.bold = True
    run1.italic = True

    run2 = para.add_run(texte_normal1)
    run2.font.size = Pt(12)
    run2.bold = False
    run2.italic = True

    if texte_gras2:
        run3 = para.add_run(texte_gras2)
        run3.font.size = Pt(12)
        run3.bold = True
        run3.italic = True

    if texte_normal3:
        run4 = para.add_run(texte_normal3)
        run4.font.size = Pt(12)
        run4.bold = False
        run4.italic = True

    if texte_gras3:
        run5 = para.add_run(texte_gras3)
        run5.font.size = Pt(12)
        run5.bold = True
        run5.italic = True

    if texte_normal4:
        run6 = para.add_run(texte_normal4)
        run6.font.size = Pt(12)
        run6.bold = False
        run6.italic = True

    if texte_gras4:
        run7 = para.add_run(texte_gras4)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
   
    if texte_normal5:
        run1 = para.add_run(texte_normal5)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True
   
    if texte_gras5:
        run7 = para.add_run(texte_gras5)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
   
    if texte_normal6:
        run1 = para.add_run(texte_normal6)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True
   
    if texte_gras6:
        run7 = para.add_run(texte_gras6)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
   
    if texte_normal7:
        run1 = para.add_run(texte_normal7)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True

def set_cell_border(cell, **kwargs):
    """
    Appliquer les bordures à une cellule de tableau.
    kwargs: top, bottom, start, end = {'sz': 12, 'val': 'single', 'color': '000000'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'bottom', 'start', 'end'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in kwargs[edge]:
                element.set(qn('w:{}'.format(key)), kwargs[edge][key])
            tcPr.append(element)

def set_cell_border1(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in borders:
            edge_data = borders[edge]
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
            edge_el.set(qn('w:sz'), str(edge_data.get('sz', 8)))
            edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
            edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

rang_percentile_corresp = {
    (0, 9): "zone faible",
   
    (10, 25): "moyenne faible",
   
    (26, 74): "moyenne",
   
    (75, 84): "moyenne forte",
   
    (85, 95): "zone élevée",
   
    (96, 100): "zone très élevée"
}

notes_stand_corresp = {
    (1, 6): "zone faible",
   
    (7, 8): "moyenne faible",
   
    (9, 11): "moyenne",
   
    (12, 13): "moyenne forte",
   
    (14, 15): "zone élevée",
   
    (16, 19): "zone très élevée"
}

rang_per_corresp = {
    1: 0.1,
   
    2: 0.4,
   
    3: 1,
   
    4: 2,
   
    5: 5,
   
    6: 9,

    7: 16,

    8: 25,

    9: 37,

    10: 50,

    11: 63,

    12: 75,

    13: 84,

    14: 91,

    15: 95,

    16: 98,

    17: 99,

    18: 99.6,

    19: 99.9
}


def get_category(rang_percentile):
    for (min_val, max_val), category in rang_percentile_corresp.items():
        if min_val <= int(rang_percentile) <= max_val:
            return category
    return "Rang percentile invalide."

def get_notes_stand(notes_stand):
    for (min_val, max_val), category in notes_stand_corresp.items():
        if min_val <= int(notes_stand) <= max_val:
            return category
    return "Rang percentile invalide."


# fonctions pour determiner le cluster

# Cluster en fonction des parents
def cluster_typique_parent(parent):
    clusters = []
    for i in range (1, len(parent) - 1):
        if parent[i] <= 54:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_leger_typique_parent(parent):
    clusters = []
    for i in range (1, len(parent) - 1):
        if parent[i] >= 55 and parent[i] <= 59:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_modere_parent(parent):
    clusters = []
    for i in range (1, len(parent) - 1):
        if parent[i] >= 60 and parent[i] <= 69:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_net_parent(parent):
    clusters = []
    for i in range (1, len(parent) - 1):
        if parent[i] >= 70 :
            clusters.append(i)
    return ([cluster for cluster in clusters])

# Cluster en fonction des enseigants
def cluster_typique_enseignant(enseignant):
    clusters = []
    for i in range (1, len(enseignant) - 1):
        if enseignant[i] <= 54:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_leger_typique_enseignant(enseignant):
    clusters = []
    for i in range (1, len(enseignant) - 1):
        if enseignant[i] >= 55 and enseignant[i] <= 59:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_modere_enseignant(enseignant):
    clusters = []
    for i in range (1, len(enseignant) - 1):
        if enseignant[i] >= 60 and enseignant[i] <= 69:
            clusters.append(i)
    return ([cluster for cluster in clusters])

def cluster_net_enseignant(enseignant):
    clusters = []
    for i in range (1, len(enseignant) - 1):
        if enseignant[i] >= 70 :
            clusters.append(i)
    return ([cluster for cluster in clusters])

def total(note):
    if note <= 54:
        return "difficultés exécutives typique"
    if note <= 59:
        return "difficultés exécutives légerement atypique"
    if note <= 69:
        return "difficultés exécutives modérement atypique"
    if note >= 70:
        return "difficultés exécutives nettement atypique"
   
def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    # Créer les 4 côtés avec "none"
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'nil')  # Pas de bordure
        tcBorders.append(edge_el)

    tcPr.append(tcBorders)

# Réduire la marge
def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Change les marges internes d'une cellule Word (en cm)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(int(val * 567)))  # 1 cm = 567 twips
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

#************************************************************
# Fonctions pour génerer chaques parties du texte
#************************************************************

def cadrePrésentation(nom, prenom, date_naiss, age, lat, date, doc):
    table = doc.add_table(rows=5, cols=2)

    # Style général (bordures visibles)
    table.style = 'Table Grid'

    # Données à insérer
    left_col = ["NOM et PRENOM", "Date de naissance", "Age au moment de l’évaluation", "Latéralité", "Date du bilan"]
    right_col = [f"{nom} {prenom}", f"{date_naiss}", f"{age}", f"{lat}", f"{date}"]

    # Remplir les cellules
    for row_idx in range(5):
        # Colonne de gauche
        cell_left = table.cell(row_idx, 0).paragraphs[0]
        run_left = cell_left.add_run(left_col[row_idx])
        if row_idx == 0:
            run_left.bold = True
            run_left.font.color.rgb = RGBColor(149, 179, 215)  # Bleu doux
        cell_left.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Colonne de droite
        cell_right = table.cell(row_idx, 1).paragraphs[0]
        run_right = cell_right.add_run(right_col[row_idx])
        if row_idx == 0:
            run_right.bold = True
            run_right.font.color.rgb = RGBColor(149, 179, 215)  # Bleu doux
        cell_right.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Supprimer toutes les bordures intérieures
    rows = len(table.rows)
    cols = len(table.columns)

    for r in range(rows):
        for c in range(cols):
            clear_cell_borders(table.cell(r, c))

    # Remettre uniquement le contour extérieur
    # Haut
    for c in range(cols):
        set_cell_border1(table.cell(0, c), top={'val': 'single', 'sz': 12, 'color': '000000'})

    # Bas
    for c in range(cols):
        set_cell_border1(table.cell(rows - 1, c), bottom={'val': 'single', 'sz': 12, 'color': '000000'})

    # Gauche
    for r in range(rows):
        set_cell_border1(table.cell(r, 0), left={'val': 'single', 'sz': 12, 'color': '000000'})

    # Droite
    for r in range(rows):
        set_cell_border1(table.cell(r, cols - 1), right={'val': 'single', 'sz': 12, 'color': '000000'})


    doc.add_paragraph()

    # Indication
    indication = doc.add_paragraph()
    IndicationBleu = indication.add_run("INDICATION, PLAINTE PRINCIPALE")
    IndicationBleu.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    indication = doc.add_paragraph()
    indication.add_run("J’ai rencontré ")
    indication.add_run(f"{nom} {prenom} ").bold = True
    indication.add_run("à la demande des parents, afin de mieux comprendre son  fonctionnement cognitif.  ")
    indication = doc.add_paragraph()
    run(f"Ce présent bilan a donc pour objectif de définir le profil cognitif et comportemental de {prenom}, afin de fournir  des axes de " \
    "travail et d'accompagnement. ", indication)

    # Source information
    source_info = doc.add_paragraph()
    SourceBleu = source_info.add_run("SOURCE D'INFORMATION")
    SourceBleu.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    source_info = doc.add_paragraph()  
    run(f"Les informations pertinentes ont été recueillies auprès des parents et de {prenom}. ", source_info)
    source_info = doc.add_paragraph()
    source_info.add_run("Tests utilisés ").bold = True
    run(": WISC-V ; TEA-Ch ; NEPSY-II ; KiTAP ; BROWN ; Young-DIVA ", source_info)


def notes_compo_principales(ENS_CV=str(19), ENS_V=str(20), ENS_RF=str(23), ENS_MDT=str(13), ENS_VT=str(13), ENS_ET=str(64),
                            NC_CV=str(98), NC_V=str(100), NC_RF=str(109), NC_MDT=str(79), NC_VT=str(80), NC_ET=str(94),
                            RP_CV=str(45), RP_V=str(50), RP_RF=str(73), RP_MDT=str(8), RP_VT=str(9), RP_ET=str(34),
                            IDC_CV="89-107", IDC_V="92-108", IDC_RF="101-116", IDC_MDT="73-90", IDC_VT="73-92", IDC_ET="88-101",
                            prenom="Giuseppe"):
    # Bilan psychométrique
    doc.add_page_break()
    bilan_psycho = doc.add_paragraph()
    bilan_psycho.add_run("BILAN PSYCHOMÉTRIQUE")
    bilan_psycho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bilan_psycho = doc.add_paragraph()
    bilan_psycho.add_run("--------------------------------------------------------------------------------------------------")
    bilan_psycho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Cadre
    table2 = doc.add_table(rows=1, cols=1)

    cell2 = table2.cell(0, 0)
    cadre2 = cell2.add_paragraph()
    cadre2.alignment = 1
    texteGras = cadre2.add_run("Fonctionnement intellectuel global (WISC V)")
    texteGras.bold = True
    # Capacité cognitives globales
    CapCo = doc.add_paragraph()
    CapCo = doc.add_paragraph()
    # Texte bleu
    texteBleu2 = CapCo.add_run("Capacités cognitives globales ")
    texteBleu2.bold = True
    texteBleu2.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu2.underline = True
    CapCo = doc.add_paragraph()
    texteWISC = CapCo.add_run("Le WISC-V est utilisé pour mesurer les habiletés générales de raisonnement des " \
        "enfants de 6 à 16 ans. Cette  évaluation fournit un score qui représente la capacité intellectuelle globale" \
        " de l’enfant (QIT), ainsi que des  scores d’indice qui mesurent les domaines suivants du fonctionnement" \
        " cognitif : compréhension verbale  (ICV), traitement visuospatial (IVS), raison fluide (IRF), mémoire de" \
        " travail (IMT) et vitesse de traitement  (IVT). ")
    texteWISC.italic = True
    CapCo = doc.add_paragraph()
    CapCo.add_run("L’évaluation intellectuelle réalisée à l’aide du WISC-V met " \
        "en évidence un profil présentant des capacités  intellectuelles hétérogènes." \
        " En effet, l’hétérogénéité significative de son profil ne nous permet pas de " \
        f"calculer  un QIT chez {prenom}. En effet, des différences statistiquement significatives " \
        "apparaissent entre plusieurs scores  d’indices. Alors, la note d’échelle totale"
        " (QIT) – qui représente les aptitudes intellectuelles globales – ne peut " \
        f" nous permettre de comprendre le fonctionnement de {prenom}. L’étude des forces " \
        "et des faiblesses est préconisée  pour mieux comprendre son profil cognitif. ")

    # Ajouter un titre centré et en italique
    titre = doc.add_paragraph("Synthèse des notes composites principales")
    doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.runs[0].italic = True
    # Add table with 7 columns and 7 rows
    table = doc.add_table(rows=7, cols=7)
    table.style = 'Table Grid'
    table.autofit = True

    # Set column widths (approximate)
    col_widths = [Inches(2.2), Inches(0.9), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.4), Inches(1.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Header row
    headers = [
        "Composite", "Indices", "Ensemble des Note Standard", "Note Composite",
        "Rang Percentile", "Intervalle de Confiance", "Description qualitative"
    ]
    for i, text in enumerate(headers):
        p = table.cell(0, i).paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    data = [
        ["Compréhension Verbale", "ICV", str(ENS_CV), str(NC_CV), str(RP_CV), str(IDC_CV), get_category(int(RP_CV))],
        ["Visuospatial", "IVS", str(ENS_V), str(NC_V), str(RP_V), str(IDC_V), get_category(int(RP_V))],
        ["Raisonnement Fluide", "IRF", str(ENS_RF), str(NC_RF), str(RP_RF), str(IDC_RF), get_category(int(RP_RF))],
        ["Mémoire de travail", "IMT", str(ENS_MDT), str(NC_MDT), str(RP_MDT), str(IDC_MDT), get_category(RP_MDT)],
        ["Vitesse de traitement", "IVT", str(ENS_VT), str(NC_VT), str(RP_VT), str(IDC_VT), get_category(RP_VT)],
        ["Échelle Totale", "QIT", str(ENS_ET), str(NC_ET), str(RP_ET), str(IDC_ET), get_category(RP_ET)]
    ]

    # Fill in data rows
    for i, row_data in enumerate(data):
        row_idx = i + 1

        # Déterminer la couleur de la ligne à partir de la colonne 4
        category = get_category(row_data[4])
        if category == "zone faible":
            color = RGBColor(255, 41, 41)
        elif category == "moyenne faible":
            color = RGBColor(255, 105, 41)
        elif category == "moyenne":
            color = RGBColor(0, 0, 0)
        elif category == "moyenne forte":
            color = RGBColor(52, 168, 83)
        elif category == "zone élevée":
            color = RGBColor(59, 115, 47)
        elif category == "zone très élevée":
            color = RGBColor(23, 78, 166)
        else:
            color = RGBColor(0, 0, 0)  # default (black)

        # Remplir les cellules de la ligne avec la couleur déterminée
        for j, val in enumerate(row_data):
            cell = table.cell(row_idx, j)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.color.rgb = color

            # Extrait spécifique
            if row_idx == 6 and j == 6:
                run.font.strike = True

        for row in table.rows:
            row.height = Pt(20)
           

    doc.add_paragraph()

#************************************************************************************************
# INDICES
#************************************************************************************************

def indices(IAG=str(15), RP1=str(95), ICC=str(13), RP2=str(84), INV=str(33), RP3=str(100), prenom="Giuseppe"):
    # Description après tableau
    ajouter_paragraphe(
        doc,
        f"L'indice complémentaire d'aptitude générale (IAG = {IAG}, RP = {RP1})",
        ", témoigne d'une capacité à raisonner,de facultés de compréhension, situées dans la ",
        get_category(RP1),
        " de son âge."
    )
    doc.add_paragraph()
    ajouter_paragraphe(
        doc,
        f"L'indice de compétence cognitive (ICC = {ICC}, RP = {RP2}) ",
        "relatif aux traitements de bas niveaux (vitesse detraitement et mémoire de travail) se situe dans la ",
        get_category(RP2),
        " de son âge."
    )
    doc.add_paragraph()
    ajouter_paragraphe(
        doc,
        f"L'indice non verbal (INV = {INV}, RP = {RP3}) ",
        "relatif aux compétences non verbales se situe dans la ",
        get_category(RP3),
        " de son âge."
    )
    doc.add_paragraph()
    ajouter_paragraphe(
        doc,
        "",
        "Nous détaillerons, dans les chapitres suivants, les différents domaines cognitifs qui ont été évaluées," \
        f" et qui permettent d'appréhender de manière plus approfondie le fonctionnement actuel de {prenom}.",
        "",
        ""
    )

#************************************************************************************************
# CAPACITE VERBAL
#************************************************************************************************

def capacite_verbal(ICV=str(28), RP=str(92), note_stand_simi=str(15), note_stand_vocab=str(13), prenom="Giuseppe"):
    # Correspondance des rangs percentiles en fonction des notes standards
    rang_per_simi = rang_per_corresp[int(note_stand_simi)]
    rang_per_vocab = rang_per_corresp[int(note_stand_vocab)]

    doc.add_page_break()
    CapVe = doc.add_paragraph()
    # Texte bleu
    # Capacités verbales
    texteBleu3 = CapVe.add_run("Capacités verbales")
    texteBleu3.bold = True
    texteBleu3.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu3.underline = True

    ajouter_paragraphe_italique(
        doc,
        "L'indice de Compréhension Verbale (ICV) ",
        "mesure les aptitudes verbales en sollicitant le raisonnement, la compréhension, et la catégorisation. Il évalue la formation de concepts verbaux et les connaissances culturelles acquises dans l'environnement de l'enfant."
    )

    ajouter_paragraphe(
        doc,
        "",
        f"Sur le plan qualitatif, nous observons que {prenom} comprend bien les consignes au décours du bilan. Il n'aura pas eu besoin d'énormément de reformulations ou d'explications supplémentaires. Le discours spontané est fluent, informatif et cohérent.",
        "",
        ""
    )

    ajouter_paragraphe(
        doc,
        "",
        "La note composite de compréhension verbale ",
        f"(ICV = {ICV}, RP = {RP}) ",
        "se situe dans la ",
        get_category(RP),
        " comparaison aux enfants du même âge."
    )

    # Tableau : Epreuve, Notes standards, Rang Percentile
    # Création du tableau
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Données
    entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
    matrices = ["Similitudes", note_stand_simi, rang_per_simi]
    balances = ["Vocabulaire", note_stand_vocab, rang_per_vocab]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)


    # Ligne 1 : Matrices
    for i, val in enumerate(matrices):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # Ligne 2 : Balances (vert + italique)
    for i, val in enumerate(balances):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)
   
    doc.add_paragraph()

    # Partie similitude
    if get_category(rang_per_simi) == "zone faible":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre des difficultés de catégorisation et de conceptualisation."
        )
   
    if get_category(rang_per_simi) == "moyenne faible":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre des fragilités dans la catégorisation et la conceptualisation."
        )

    if get_category(rang_per_simi) == "moyenne":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre de bonnes capacités de catégorisation et de conceptualisation."
        )

    if get_category(rang_per_simi) == "moyenne forte":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre de bonnes capacités de catégorisation et de conceptualisation."
        )

    if get_category(rang_per_simi) == "zone elevée":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre de très bonnes capacités de catégorisation et de conceptualisation."
        )

    if get_category(rang_per_simi) == "zone très elevée":
        ajouter_paragraphe(
        doc,
        "",
        "Le subtest « ",
        "Similitudes ",
        f"» permet d'appréhender la qualité du langage oral de {prenom}, ses capacités d'abstractionet de raisonnement verbal. {prenom} présente des performances dans la ",
        get_category(rang_per_simi),
        " de son âge. Il montre d’excellentes capacités de catégorisation et de conceptualisation."
        )

    ajouter_paragraphe_italique(
        doc,
        "Exemple : « En quoi le cochon et le mouton se ressemblent ? Qu’est-ce qui fait qu’ils sont pareil ? »",
        ""
    )

    # Partie vocabulaire

    if get_category(rang_per_vocab) == "zone faible":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un stock lexical pauvre."
        )
   
    if get_category(rang_per_vocab) == "moyenne faible":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un stock lexical fragile."
        )

    if get_category(rang_per_vocab) == "moyenne":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un bon stock lexical."
        )
   
    if get_category(rang_per_vocab) == "moyenne forte":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un bon stock lexical."
        )

    if get_category(rang_per_vocab) == "zone élevée":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un très bon stock lexical."
        )

    if get_category(rang_per_vocab) == "zone très élevée":
        ajouter_paragraphe(
        doc,
        "",
        "Lors de l'épreuve de «",
        "Vocabulaire",
        f"», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), {prenom} obtient des résultats dans la ",
        get_category(rang_per_vocab),
        f" de son âge. {prenom} témoigne d’un excellent stock lexical."
        )

    ajouter_paragraphe_italique(
        doc,
        "Exemple : « Qu'est-ce qu'une fourchette ? »",
        ""
    )
    doc.add_paragraph()
   
#************************************************************************************************
# FONCTION VISUO-SPATIAL
#************************************************************************************************

def visuo_spatial(IVS=str(30), RP=str(97), note_stand_cube=str(16), note_stand_puzz=str(14), prenom="Giuseppe"):
    # Correspondance des rangs percentiles en fonction des notes standards
    rang_per_cube = rang_per_corresp[int(float(note_stand_cube))]
    rang_per_puzz = rang_per_corresp[int(float(note_stand_puzz))]

    # Fonctions visuo-spatiales
    # Texte bleu
    doc.add_page_break()
    FVS = doc.add_paragraph()
    texteBleu4 = FVS.add_run("Fonctions visuo-spatiales")
    texteBleu4.bold = True
    texteBleu4.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu4.underline = True

    ajouter_paragraphe_italique(
        doc,
        "L'indice Visuo-Spatial (IVS)",
        ", mesure la capacité à analyser les détails visuels et comprendre les relations  visuo-spatiales afin de construire des dessins géométriques à partir d'un modèle. Cette habileté requiert un  raisonnement visuo-spatial, l’intégration et la synthèse de relations « partie-tout », l'attention aux détails  visuels et l'intégration visuo-motrice. "
    )
    doc.add_paragraph()
    ajouter_paragraphe(
        doc,
        "",
        "Dans ce domaine ",
        f"(IVS = {str(IVS)}, RP = {str(RP)}) ",
        f"{prenom} possède des capacités visuo-constructives, d’analyse visuo spatiale, et de résolution de problème dans la ",
        get_category(RP),
        " de son âge."
    )

    # Tableau Epreuve cubes puzzle
    # Création du tableau
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Données
    entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
    matrices = ["Cubes", note_stand_cube, rang_per_cube]
    balances = ["Puzzles visuels", note_stand_puzz, rang_per_puzz]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # Ligne 1 : Matrices
    for i, val in enumerate(matrices):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # Ligne 2 : Balances (vert + italique)
    for i, val in enumerate(balances):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # === Texte + Image cubes
    doc.add_paragraph()
    row = doc.add_table(rows=1, cols=2)
    row.autofit = False
    table.columns[1].width = Cm(50)

    # Texte à gauche
    col_text = row.cell(0, 0)
    paragraph = col_text.paragraphs[0]
    paragraph.add_run("Le subtest « ").italic = True
    paragraph.add_run("Cubes").bold = True
    paragraph.add_run(" », où il est demandé à Léo de reproduire des patterns visuels à l’aide de cubes bicolores, Léo obtient des résultats dans la ").italic = True
    paragraph.add_run(get_category(rang_per_cube)).bold = True
    if get_category(rang_per_cube) == "zone faible":
        paragraph.add_run(f" de son âge. {prenom} montre des difficultés dans l’organisation spatiale des modèles et dans la visuo-construction.")
   
    if get_category(rang_per_cube) == "moyenne faible":
        paragraph.add_run(f" de son âge. {prenom} présente des fragilités dans l’organisation spatiale des modèles et la visuo-construction.")
   
    if get_category(rang_per_cube) == "moyenne":
        paragraph.add_run(f" de son âge. {prenom} montre de bonnes capacités dans l’organisation spatiale des modèles et dans la visuo-construction.")
   
    if get_category(rang_per_cube) == "moyenne forte":
        paragraph.add_run(f" de son âge. {prenom} montre de bonnes capacités dans l’organisation spatiale des modèles et dans la visuo-construction.")
   
    if get_category(rang_per_cube) == "zone élevée":
        paragraph.add_run(f" de son âge. {prenom} montre de très bonnes capacités dans l’organisation spatiale des modèles et dans la visuo-construction.")
   
    if get_category(rang_per_cube) == "zone très élevée":
        paragraph.add_run(f" de son âge. {prenom} montre d’excellentes capacités dans l’organisation spatiale des modèles et dans la visuo-construction.")

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Image à droite
    col_img = row.cell(0, 1)
    run_img = col_img.paragraphs[0].add_run()
    run_img.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/cube.png", width=Inches(1.5))
    col_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER    


    # Paragraphe descriptif "Puzzles visuels"

    # === Texte + Image cubes
    doc.add_paragraph()
    row2 = doc.add_table(rows=1, cols=2)
    row2.autofit = False

    # Texte à gauche
    col_text2 = row2.cell(0, 0)
    paragraph2 = col_text2.paragraphs[0]
    paragraph2.add_run("Pour le subtest « ").italic = True
    paragraph2.add_run("Puzzles visuels").bold = True
    paragraph2.add_run(" », où il lui est demandé de choisir trois piècesde puzzle qui, ensemble, reconstruiraient le modèle visuel, Léo présente des performances dans la ").italic = True
    paragraph2.add_run(f"{prenom} se situe dans la ").italic = True
    paragraph2.add_run(get_category(rang_per_puzz)).bold = True
    if get_category(rang_per_puzz) == "zone faible":
        paragraph2.add_run(f" de son âge. {prenom} présente des difficultés d’analyse visuo-spatiale.")

    if get_category(rang_per_puzz) == "moyenne faible":
        paragraph2.add_run(f" de son âge. {prenom} présente des fragilités dans l’analyse visuo-spatiale.")
   
    if get_category(rang_per_puzz) == "moyenne":
            paragraph2.add_run(f" de son âge. {prenom} présente de bonnes compétences dans l’analyse visuo-spatiale.")
   
    if get_category(rang_per_puzz) == "moyenne forte":
            paragraph2.add_run(f" de son âge. {prenom} présente de bonnes compétences dans l’analyse visuo-spatiale.")

    if get_category(rang_per_puzz) == "zone élevée":
            paragraph2.add_run(f" de son âge. {prenom} présente de très bonnes compétences dans l’analyse visuo-spatiale.")
   
    if get_category(rang_per_puzz) == "zone très élevée":
            paragraph2.add_run(f" de son âge. {prenom} présente d'excellentes compétences dans l’analyse visuo-spatiale.")

    paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Image à droite
    col_img2 = row2.cell(0, 1)
    run_img2 = col_img2.paragraphs[0].add_run()
    run_img2.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/puzzles.png", width=Inches(2))
    col_img2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT




#************************************************************************************************
#RAISONNEMENT FLUIDE
#************************************************************************************************


def raisonnement_fluide(IRF=str(23), RP=str(73), note_stand_mat=str(12), note_stand_bal=str(11), prenom="Giuseppe"):
    # Correspondance des rangs percentiles en fonction des notes standards
    rang_per_mat = rang_per_corresp[int(float(note_stand_mat))]
    rang_per_bal = rang_per_corresp[int(float(note_stand_bal))]

    # Titre "Raisonnement fluide"
    doc.add_page_break()
    RF = doc.add_paragraph()
    texteBleu5 = RF.add_run("Raisonnement fluide")
    texteBleu5.bold = True
    texteBleu5.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu5.underline = True

    # Paragraphe explicatif en italique
    ajouter_paragraphe_italique(
        doc,
        "L’indice de raisonnement fluide (IRF), ",
        f"permet de mesurer la capacité de {prenom} à détecter la relation conceptuelle sous-jacente entre des images et à utiliser le raisonnement pour identifier et appliquer des règles. L’identification et l’application des relations conceptuelles dans l’IRF exigent un raisonnement inductif et  quantitatif, l’intelligence à grande échelle, le traitement simultané et la pensée abstraite. "
    )

    # Texte général IRF
    ajouter_paragraphe(
        doc,
        "",
        "Les capacités de raisonnement fluides",
        f" (IRF = {IRF}, RP = {RP})",
        " telles que mesurées par le WISC-V, apparaissent ce jour, dans la ",
        get_category(RP),
        " pour son âge."
    )

    doc.add_paragraph()

    # Création du tableau
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Données
    entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
    matrices = ["Matrices", note_stand_mat, rang_per_mat]
    balances = ["Balances", note_stand_bal, rang_per_bal]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 1 : Matrices
    for i, val in enumerate(matrices):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 2 : Balances (orange)
    for i, val in enumerate(balances):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # === Texte + Image cubes
    doc.add_paragraph()
    row = doc.add_table(rows=1, cols=2)
    row.autofit = False

    # Texte à gauche
    col_text = row.cell(0, 0)
    paragraph = col_text.paragraphs[0]
    paragraph.add_run("Le subtest « ").italic = True
    paragraph.add_run("Matrices").bold = True
    paragraph.add_run(" » sollicite les compétences visuo-spatiales, le raisonnement visuel, les capacités d’induction (inférer une logique à partir de l’observation), les capacités de déduction (généralisation d’une logique et application sur de nouveaux éléments) ainsi qu’une démarche catégorielle (abstraction de traits communs et de différences par comparaisons). Ces derniers se situent dans la ").italic = True
    paragraph.add_run(get_category(rang_per_mat)).bold = True
    if get_category(rang_per_mat) == "zone faible":
            paragraph.add_run(f" de son âge. {prenom} montre des difficultés en logique d’ordre visuo-spatiale. ")

    if get_category(rang_per_mat) == "moyenne faible":
            paragraph.add_run(f" de son âge. {prenom} montre des fragilités dans la logique d’ordre visuo-spatiale. ")
   
    if get_category(rang_per_mat) == "moyenne":
            paragraph.add_run(f" de son âge. {prenom} montre de bonnes aptitudes en logique d’ordre visuo-spatiale. ")

    if get_category(rang_per_mat) == "moyenne forte":
            paragraph.add_run(f" de son âge. {prenom} montre de bonnes aptitudes en logique d’ordre visuo-spatiale. ")

    if get_category(rang_per_mat) == "zone élevée":
            paragraph.add_run(f" de son âge. {prenom} montre de très bonnes aptitudes en logique d’ordre visuo-spatiale. ")

    if get_category(rang_per_mat) == "zone très élevée":
            paragraph.add_run(f" de son âge. {prenom} montre d’excellentes aptitudes en logique d’ordre visuo-spatiale. ")


    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Image à droite
    col_img = row.cell(0, 1)
    run_img = col_img.paragraphs[0].add_run()
    run_img.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/matrices.png", width=Inches(2))
    col_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER    


    # === Texte + Image cubes
    doc.add_paragraph()
    row2 = doc.add_table(rows=1, cols=2)
    row2.alignment = WD_TABLE_ALIGNMENT.CENTER
    row2.autofit = False

    row2.columns[0].width = Cm(40)
    row2.columns[1].width = Cm(5)

    # Texte à gauche
    col_text2 = row2.cell(0, 0)
    paragraph2 = col_text2.paragraphs[0]
    paragraph2.add_run("Pour le subtest « ").italic = True
    paragraph2.add_run("Balances").bold = True
    paragraph2.add_run(" », il s’agit d’une tâche de logique inductive et déductive pour laquelle le concept quantitatif d’égalité doit être acquis afin de permettre l’application des concepts de correspondance, d’addition et/ou de multiplication. Dans cette épreuve, ").italic = True
    paragraph2.add_run(f"{prenom} se situe dans la ").italic = True
    paragraph2.add_run(get_category(rang_per_bal)).bold = True

    if get_category(rang_per_bal) == "zone faible":
            paragraph2.add_run(f" de son âge. {prenom} présente des difficultés dans le raisonnement logico-mathématique.")
   
    if get_category(rang_per_bal) == "moyenne faible":
            paragraph2.add_run(f" de son âge. {prenom} présente des fragilités dans le raisonnement logico-mathématique.")

    if get_category(rang_per_bal) == "moyenne":
            paragraph2.add_run(f" de son âge. {prenom} montre un bon raisonnement logico-mathématique.")

    if get_category(rang_per_bal) == "moyenne forte":
            paragraph2.add_run(f" de son âge. {prenom} montre un bon raisonnement logico-mathématique.")

    if get_category(rang_per_bal) == "zone élevée":
        paragraph2.add_run(f" de son âge. {prenom} montre un très bon raisonnement logico-mathématique.")

    if get_category(rang_per_bal) == "zone très élevée":
            paragraph2.add_run(f" de son âge. {prenom} montre un très bon raisonnement logico-mathématique.")

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Image à droite
    col_img2 = row2.cell(0, 1)
    run_img2 = col_img2.paragraphs[0].add_run()
    run_img2.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/balances.png", width=Inches(2))
    col_img2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


#************************************************************************************************
# MEMOIRE DE TRAVAIL
#************************************************************************************************

def memoire_de_travail(IMT=str(19), RP=str(42), note_stand_chiffre=str(10), note_stand_image=str(9), prenom="Giuseppe"):
    # Correspondance des rangs percentiles en fonction des notes standards
    rang_per_chiffre = rang_per_corresp[int(float(note_stand_chiffre))]
    rang_per_image = rang_per_corresp[int(float(note_stand_image))]

    # Texte bleu
    doc.add_page_break()
    MDT = doc.add_paragraph()
    texteBleu6 = MDT.add_run("Mémoire de travail")
    texteBleu6.bold = True
    texteBleu6.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu6.underline = True

    ajouter_paragraphe_italique(
        doc,
        "La mémoire de travail (IMT) ",
        "est la capacité à manipuler de l’information que l’on maintient en mémoire à  court terme (ex. garder en tête les étapes d’un calcul mental, tout en effectuant la tâche de calcul). La mémoire à court terme représente un ensemble de processus qui permet la gestion d’un flux d’informations,  et leur stockage temporaire. Cette mémoire correspond à la quantité d’informations que l’on peut maintenir  active (ex. retenir un numéro de téléphone avant de le noter sur un papier).  "
    )

    doc.add_paragraph()

    ajouter_paragraphe(
        doc,
        "",
        "Les résultats à cet indice (",
        f"IMT = {IMT}, RP = {RP}",
        ") témoigne d’une mémoire de travail dans la ",
        get_category(RP),
        " de son âge"
    )

    doc.add_paragraph()

    # Tableau memoire des chiffres et des images
    # Création du tableau
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Données
    entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
    matrices = ["Mémoire des Chiffres ", note_stand_chiffre, rang_per_chiffre]
    balances = ["Mémoire des Images ", note_stand_image, rang_per_image]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 1 : Matrices
    for i, val in enumerate(matrices):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 2 : Balances (orange)
    for i, val in enumerate(balances):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    ajouter_paragraphe_italique(
        doc,
        "",
        "Le versant auditivo-verbale met en jeu la boucle phonologique afin de maintenir et manipuler les informations  auditivo-verbales pour la réalisation d’une tâche. Le versant visuo-spatiale met en jeu le calepin visuo-spatial  afin de maintenir et manipuler des informations visuelles pour réaliser la tâche. "
    )

    doc.add_paragraph()

    # === Texte + Image cubes
    doc.add_paragraph()
    row = doc.add_table(rows=1, cols=2)
    row.autofit = False

    # Texte à gauche
    col_text = row.cell(0, 0)
    paragraph = col_text.paragraphs[0]
    paragraph.add_run("Lors du subtest « ").italic = True
    paragraph.add_run("Mémoire des chiffres").bold = True
    paragraph.add_run(f" », {prenom} présente des performances dans la ").italic = True
    paragraph.add_run(get_category(rang_per_chiffre)).bold = True

    if get_category(rang_per_chiffre) == "zone faible":
        paragraph.add_run(f" de son âge. {prenom} présente des difficultés en mémoire à court terme sur un support auditif.")

    if get_category(rang_per_chiffre) == "moyenne faible":
        paragraph.add_run(f" de son âge. {prenom} présente de légères difficulté en mémoire à court terme sur un support auditif.")

    if get_category(rang_per_chiffre) == "moyenne":
        paragraph.add_run(f" de son âge. {prenom} présente de bonnes compétences en mémoire à court terme sur un support auditif.")

    if get_category(rang_per_chiffre) == "moyenne forte":
        paragraph.add_run(f" de son âge. {prenom} présente de bonnes compétences en mémoire à court terme sur un support auditif.")

    if get_category(rang_per_chiffre) == "zone élevée":
        paragraph.add_run(f" de son âge. {prenom} présente de très bonnes compétences en mémoire à court terme sur un support auditif.")

    if get_category(rang_per_chiffre) == "zone très élevée":
        paragraph.add_run(f" de son âge.  {prenom} présente de très bonnes compétences en mémoire à court terme sur un support auditif.")

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Image à droite
    col_img = row.cell(0, 1)
    run_img = col_img.paragraphs[0].add_run()
    run_img.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/memoire_des_chiffres.png", width=Inches(2))
    col_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    # === Texte + Image cubes
    doc.add_paragraph()
    row2 = doc.add_table(rows=1, cols=2)
    row2.autofit = False

    # Texte à gauche
    col_text2 = row2.cell(0, 0)
    paragraph2 = col_text2.paragraphs[0]
    paragraph2.add_run("Lors du subtest « ").italic = True
    paragraph2.add_run("Mémoire des images").bold = True
    paragraph2.add_run(f" », {prenom} se situe dans la ").italic = True
    paragraph2.add_run(get_category(rang_per_image)).bold = True

    if get_category(rang_per_image) == "zone faible":
        paragraph.add_run(f" de son âge. {prenom} présente des difficultés en mémoire à court terme sur un support visuel. Lors de cette épreuve, il arrive vite à saturation et  ne semble plus disposer des ressources attentionnelles nécessaire pour continuer l’épreuve.")

    if get_category(rang_per_image) == "moyenne faible":
        paragraph2.add_run(f" de son âge. {prenom} présente des fragilités en mémoire à court terme sur un support visuel.")

    if get_category(rang_per_image) == "moyenne":
        paragraph2.add_run(f" de son âge. {prenom} présente de bonnes capacités en mémoire à court terme sur un support visuel.")

    if get_category(rang_per_image) == "moyenne forte":
        paragraph2.add_run(f" de son âge. {prenom} présente de bonnes capacités en mémoire à court terme sur un support visuel.")

    if get_category(rang_per_image) == "zone élevée":
        paragraph2.add_run(f" de son âge. {prenom} présente de très bonnes capacités en mémoire à court terme sur un support visuel.")

    if get_category(rang_per_image) == "zone très élevée":
        paragraph2.add_run(f" de son âge. {prenom} présente d’excellentes bonnes capacités en mémoire à court terme sur un support visuel.")

       
    # Image à droite
    col_img2 = row2.cell(0, 1)
    run_img2 = col_img2.paragraphs[0].add_run()
    run_img2.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/memoire_des_images.png", width=Inches(2))
    col_img2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


#************************************************************************************************
# VITESSE DE TRAITEMENT
#************************************************************************************************

def vitesse_de_traitement(IVT=str(32), RP=str(99), note_stand_code=str(15), note_stand_symb=str(17), prenom="Giuseppe"):
    # Correspondance des rangs percentiles en fonction des notes standards
    rang_per_code = rang_per_corresp[int(float(note_stand_code))]
    rang_per_symb = rang_per_corresp[int(float(note_stand_symb))]

    # Vitesse de traitement
    doc.add_page_break()
    VDT = doc.add_paragraph()
    texteBleu6 = VDT.add_run("Vitesse de traitement")
    texteBleu6.bold = True
    texteBleu6.font.color.rgb = RGBColor(121, 181, 237)  # Couleur
    texteBleu6.underline = True

    ajouter_paragraphe_italique(
        doc,
        "",
        "La vitesse de traitement de l’information (IVT) réfère au rythme auquel un individu passe à travers différentes  étapes de cognition lors de l’exécution d’une tâche. Ces subtests font appel à la discrimination visuelle, à  l’organisation perceptive mais aussi aux habiletés graphiques. "
    )

    ajouter_paragraphe(
        doc,
        "",
        "La vitesse de traitement (",
        f"IVT = {IVT}, RP = {RP}",
        "), telle que mesurée par le WISC-V, apparaissent ce jour dans la ",
        get_category(RP),
        " de son âge."
    )

    # Tableau code et symbole
    # Création du tableau
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Données
    entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
    matrices = ["Code", note_stand_code, rang_per_code]
    balances = ["Symboles", note_stand_symb, rang_per_symb]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 1 : Matrices
    for i, val in enumerate(matrices):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 2 : Balances (orange)
    for i, val in enumerate(balances):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        if i == 1:
            if get_notes_stand(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_notes_stand(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_notes_stand(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_notes_stand(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_notes_stand(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_notes_stand(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)

        if i == 2:
            if get_category(val) == "zone faible":
                run.font.color.rgb = RGBColor(255, 41, 41)
            if get_category(val) == "moyenne faible":
                run.font.color.rgb = RGBColor(255, 105, 41)
            if get_category(val) == "moyenne":
                run.font.color.rgb = RGBColor(0, 0, 0)
            if get_category(val) == "moyenne forte":
                run.font.color.rgb = RGBColor(52, 168, 83)
            if get_category(val) == "zone élevée":
                run.font.color.rgb = RGBColor(59, 115, 47)
            if get_category(val) == "zone très élevée":
                run.font.color.rgb = RGBColor(23, 78, 166)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

        # === Texte + Image cubes
    doc.add_paragraph()
    row = doc.add_table(rows=1, cols=2)
    row.autofit = False

    # Texte à gauche
    col_text = row.cell(0, 0)
    paragraph = col_text.paragraphs[0]
    paragraph.add_run("Lors du subtest « ").italic = True
    paragraph.add_run("Code").bold = True
    paragraph.add_run(f" » {prenom} présente des performances dans la ").italic = True
    paragraph.add_run(get_category(rang_per_code)).bold = True

    if get_category(rang_per_code) == "zone faible":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve de difficultés dans la reconnaissance et la discrimination visuelle fine et montre une vitesse de traitement lente. ")

    if get_category(rang_per_code) == "moyenne faible":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve de fragilités dans la reconnaissance et la discrimination visuelle fine et montre une vitesse de traitement ralentie.")

    if get_category(rang_per_code) == "moyenne":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve de bonnes capacités dans la reconnaissance et la discrimination visuelle fine et montre une vitesse de traitement satisfaisante.")

    if get_category(rang_per_code) == "moyenne forte":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve de bonnes capacités dans la reconnaissance et la discrimination visuelle fine et montre une bonne vitesse de traitement. ")

    if get_category(rang_per_code) == "zone élevée":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve de très bonnes capacités dans la reconnaissance et la discrimination visuelle fine avec une vitesse de traitement plus que satisfaisante. ")

    if get_category(rang_per_code) == "zone très élevée":
        paragraph.add_run(f" comparativement aux enfants de son âge. {prenom} fait preuve d’excellentes capacités dans la reconnaissance et la discrimination visuelle fine fine avec une vitesse de traitement plus que satisfaisante. ")

    # Image à droite
    col_img = row.cell(0, 1)
    run_img = col_img.paragraphs[0].add_run()
    run_img.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/codes.png", width=Inches(2))
    col_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


    # === Texte + Image cubes
    doc.add_paragraph()
    row2 = doc.add_table(rows=1, cols=2)
    row2.autofit = False

    # Texte à gauche
    col_text2 = row2.cell(0, 0)
    paragraph2 = col_text2.paragraphs[0]
    paragraph2.add_run("Lors du subtest « ").italic = True
    paragraph2.add_run("Symboles").bold = True
    paragraph2.add_run(f" », {prenom} se situe dans la ").italic = True
    paragraph2.add_run(get_category(rang_per_symb)).bold = True

    if get_category(rang_per_symb) == "zone faible":
        paragraph2.add_run(f" de son âge, suggérant une vitesse de traitement lente lors de ce subtest.")

    if get_category(rang_per_symb) == "moyenne faible":
        paragraph2.add_run(f" de son âge, suggérant une certaine lenteur dans la vitesse de traitement lors de ce subtest.")

    if get_category(rang_per_symb) == "moyenne":
        paragraph2.add_run(f" de son âge, suggérant une bonne vitesse de traitement lors de ce subtest.")

    if get_category(rang_per_symb) == "moyenne forte":
        paragraph2.add_run(f" de son âge, suggérant une bonne vitesse de traitement lors de ce subtest.")

    if get_category(rang_per_symb) == "zone élevée":
        paragraph2.add_run(f" de son âge, suggérant une très bonne vitesse de traitement lors de ce subtest.")
   
    if get_category(rang_per_symb) == "zone très élevée":
        paragraph2.add_run(f" de son âge, suggérant une excellente vitesse de traitement lors de ce subtest.")

    # Image à droite
    col_img2 = row2.cell(0, 1)
    run_img2 = col_img2.paragraphs[0].add_run()
    run_img2.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/symboles.png", width=Inches(2))
    col_img2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)


    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(6)




    doc.add_paragraph()

#************************************************************************************************
# CONCLUSION
#************************************************************************************************

def conclusion(nom, prenom):
    # Conclusion
    paraConclusion = doc.add_paragraph()
    runCclusion = paraConclusion.add_run("CONCLUSION")
    runCclusion.bold = True
    runCclusion.font.size = Pt(15)
    paraConclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    ajouter_paragraphe(
        doc,
        "",
        f"L’évaluation psychométrique de {prenom} met en évidence ",
        "de bonnes capacités intellectuelles",
        ", notamment dans  les domaines du raisonnement verbal, visuo-spatial et du raisonnement fluide, qui se situent dans la moyenne  de son âge. Ses compétences verbales sont satisfaisantes, son raisonnement fluide ainsi que ses capacités  d’analyse visuo-spatiale sont également bien développés. Ces résultats témoignent d’un potentiel cognitif  préservé, malgré un profil globalement hétérogène. "
    )

    ajouter_paragraphe(
        doc,
        "",
        "Toutefois, les conditions d’évaluation ont été marquées par ",
        "une grande agitation motrice et une attention  très fluctuante, avec de nombreux décrochages ",
        f"tout au long des épreuves. {prenom} s’est montré ",
        "distractible  par l’environnement",
        ", avec de fréquentes digressions verbales, rendant certaines tâches difficiles à mener  jusqu’à leur terme. Lors des épreuves plus coûteuses sur le plan attentionnel, telles que celles impliquant la  mémoire de travail ou la vitesse de traitement, ",
        "une saturation rapide ",
        "a été observée, traduisant une difficulté  à maintenir l’effort cognitif " \
        "dans la durée. Ces éléments ont pu impacter négativement certaines de ses  " \
        "performances, en particulier dans les domaines sollicitant les ressources " \
        "attentionnelles soutenues et les capacités de concentration. "
    )

    doc.add_paragraph()

    ajouter_paragraphe(
        doc,
        "",
        "En somme, les résultats de ce bilan suggèrent un ",
        "fonctionnement cognitif globalement satisfaisant ",
        "mais  entravé par des ",
        "difficultés majeures dans la régulation attentionnelle."
        )
   
#************************************************************************************************
# INTERPREATION – ECHELLES BROWN EF/A
#************************************************************************************************

def interpretation(parent_activ, parent_focus, parent_effort, parent_emotion, parent_memoire, parent_action, parent_total,
                   enseign_activ, enseign_focus, enseign_effort, enseign_emotion, enseign__memoire, enseign_action, enseign_total,
                   nom, prenom):
    # INTERPREATION – ECHELLES BROWN EF/A
    paraConclusion = doc.add_paragraph()
    runCclusion = paraConclusion.add_run("INTERPREATION – ECHELLES BROWN EF/A")
    runCclusion.bold = True
    runCclusion.font.size = Pt(15)
    paraConclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ajouter_paragraphe(
        doc,
        "",
        "Les échelles Brown EF/A fournissent une Note Totale et six Notes Cluster. La Note Totale est la plus  exhaustive, puisqu’elle comprend tous les items du test et propose une idée générale du fonctionnement  exécutif. Les six Notes Cluster rendent comptent d’aspects plus spécifiques des fonctions exécutives et offrent  un état des lieux plus ciblés des troubles du patient. ",
        ""
    )

    doc.add_paragraph()

    ajouter_paragraphe_italique(
        doc,
        "Le Cluster 1. Activation",
        ", correspond à la capacité d’organiser, de prioriser et à démarrer le travail. ",
        "Le Cluster 2. Focus",
        ", correspond à la capacité à soutenir l’attention et à se concentrer sur des tâches. ",
        "Le Cluster 3. Effort",
        ", correspond à la capacité à maintenir l’énergie, l’effort et à ajuster la vitesse de  traitement. ",
        "Le Cluster 4. Émotion",
        ", correspond à la capacité à gérer la frustration et à moduler les émotions. ",
        " Le Cluster 5. Mémoire",
        ", correspond à la capacité à utiliser la mémoire de travail et à accéder aux  connaissances apprises.  ",
        "Le Cluster 6. Action",
        ", correspond à la capacité à surveiller et à autoréguler l’action. "
    )

    # Tableau : Cluster
    # Création du tableau : 4 lignes, 1 colonne
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'
    table.autofit = False  # Désactive le redimensionnement automatique

    # Contenu des lignes
    lignes = [
        ("70 et plus", " Nettement atypique (problème très significatif)"),
        ("60 – 69", " Modérément atypique (problème significatif)"),
        ("55-59", " Légèrement atypique (problème potentiellement important)"),
        ("54 et moins", " Typique (problème peu significatif)")
    ]

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, phrase in enumerate(lignes):
        cell = table.cell(i, 0)
        para = cell.paragraphs[0]
        run = para.add_run(phrase)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # Tableau : Activation, Focus, etc
    table = doc.add_table(rows=3, cols=8)
    table.style = 'Table Grid'
    table.autofit = False  # Désactive le redimensionnement automatique

    # Données
    entetes = ["", "Activation", "Focus", "Effort", "Émotion","Mémoire","Action","Total"]
    parent = ["Parents", parent_activ, parent_focus, parent_effort, parent_emotion, parent_memoire, parent_action, parent_total]
    enseignant = ["Enseignant(e)", enseign_activ, enseign_focus, enseign_effort, enseign_emotion, enseign__memoire, enseign_action, enseign_total]

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 1 : Matrices
    for i, val in enumerate(parent):
        cell = table.cell(1, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    # Ligne 2 : Balances (orange)
    for i, val in enumerate(enseignant):
        cell = table.cell(2, i)
        para = cell.paragraphs[0]
        run = para.add_run(str(val))
        run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

    # tableau contenant les clusters "interressant" parent
    tab_nettement_atypique_parent = [entetes[elt] for elt in cluster_net_parent(parent)]
    tab_atypique_parent = [entetes[elt] for elt in cluster_modere_parent(parent)]
    tab_legerement_atypique_parent = [entetes[elt] for elt in cluster_leger_typique_parent(parent)]
    tab_typique_parent = [entetes[elt] for elt in cluster_typique_parent(parent)]


    # tableau contenant les clusters "interressant" enseignant
    tab_nettement_atypique_enseignant = [entetes[elt] for elt in cluster_net_parent(enseignant)]
    tab_atypique_enseignant = [entetes[elt] for elt in cluster_modere_parent(enseignant)]
    tab_legerement_atypique_enseignant = [entetes[elt] for elt in cluster_leger_typique_parent(enseignant)]
    tab_typique_enseignant = [entetes[elt] for elt in cluster_typique_parent(enseignant)]


# Construction de phrase préfaite pour l'introduire dans le document
# Partie parent
    typique_parent = ""

    for elt in tab_typique_parent:
        typique_parent += elt + ", "
   
    legerement_atypique_parent = ""

    for elt in tab_legerement_atypique_parent:
        legerement_atypique_parent += elt + ", "
   
    atypique_parent = ""

    for elt in tab_atypique_parent:
        atypique_parent += elt + ", "

    nettement_atypique_parent = ""

    for elt in tab_nettement_atypique_parent:
        nettement_atypique_parent += elt + ", "

# Partie enseignant
    typique_enseignant = ""

    for elt in tab_typique_enseignant:
        typique_enseignant += elt + ", "
   
    legerement_atypique_enseignant = ""

    for elt in tab_legerement_atypique_enseignant:
        legerement_atypique_enseignant += elt + ", "
   
    atypique_enseignant = ""

    for elt in tab_atypique_enseignant:
        atypique_enseignant += elt + ", "

    nettement_atypique_enseignant = ""

    for elt in tab_nettement_atypique_enseignant:
        nettement_atypique_enseignant += elt + ", "



# Auto-completion de la partie parent
    ajouter_paragraphe(
        doc,
        "",
        "Les ",
        "parents ",
        f"de {prenom} rapportent des difficultés :  "
    )

    if tab_typique_parent != []:
        ajouter_paragraphe(
        doc,
        "- Typiques ",
        f"dans les clusters : {typique_parent} ",
        ""
        )

    if tab_legerement_atypique_parent != []:
        ajouter_paragraphe(
        doc,
        "- Legerement atypiques",
        f"dans les clusters : {legerement_atypique_parent} ",
        ""
        )

    if tab_atypique_parent != []:
        ajouter_paragraphe(
        doc,
        "- Modérément atypiques ",
        f"dans les clusters : {atypique_parent} ",
        ""
        )

    if tab_nettement_atypique_parent != []:
        ajouter_paragraphe(
        doc,
        "- Nettement atypiques ",
        f"dans les clusters : {nettement_atypique_parent} ",
        ""
        )

    ajouter_paragraphe(
        doc,
        "",
        "Au total, ces résultats mettent en avant des ",
        f"{total(parent[-1])}"
    )



# Auto-completion de la partie enseignant
    ajouter_paragraphe(
        doc,
        "",
        "L'",
        f"enseignant(e) ",
        f"de {prenom} rapenseignantt des difficultés :  "
    )


    if tab_typique_enseignant != []:
        ajouter_paragraphe(
        doc,
        "- Typiques ",
        f"dans les clusters : {typique_enseignant} ",
        ""
        )

    if tab_legerement_atypique_enseignant != []:
        ajouter_paragraphe(
        doc,
        "- Legerement atypiques",
        f"dans les clusters : {legerement_atypique_enseignant} ",
        ""
        )

    if tab_atypique_enseignant != []:
        ajouter_paragraphe(
        doc,
        "- Modérément atypiques ",
        f"dans les clusters : {atypique_enseignant} ",
        ""
        )

    if tab_nettement_atypique_enseignant != []:
        ajouter_paragraphe(
        doc,
        "- Nettement atypiques ",
        f"dans les clusters : {nettement_atypique_enseignant} ",
        ""
        )

    ajouter_paragraphe(
        doc,
        "",
        "Au total, ces résultats mettent en avant des ",
        f"{total(enseignant[-1])}"
    )


    # IMAGE
    para_img = doc.add_paragraph()
    run_img = para_img.add_run()
    run_img.add_picture("/Users/Arcimboldo/Desktop/logiciel/images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

#************************************************************************************************
# EVALUATION MANIFESTATION TDA/H
#************************************************************************************************

def evaluation_manifestation(nom, prenom):
    # EVALUATION MANIFESTATION TDA/H
    EMT = doc.add_paragraph()
    run = EMT.add_run("EVALUATION MANIFESTATION TDA/H ")
    run.bold = True
    run.font.size = Pt(15)
    EMT.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ajouter_paragraphe(
        doc,
        "",
        f"Au regard des difficultés rapportées lors de l’anamnèse, il nous est apparu nécessaire de réaliser un entretien  auprès de {prenom}. Il est important de rappeler qu’il s’agit d’un entretien clinique et que les résultats du bilan  puissent être quelque peu différents. En effet, il est possible que {prenom} soit dans de meilleures conditions lors  de l’évaluation : en dualité, sans bruit parasite ni de passage environnant.  ",
        ""
    )

    doc.add_paragraph()

    # Young Diva
    Young = doc.add_paragraph()
    run = Young.add_run("Young -DIVA 5.0 ")
    run.bold = True
    run.underline = True
    run2 = Young.add_run(": Au niveau des critères selon le DMS-5 ")

    ajouter_paragraphe(
        doc,
        "Critère A : Il y aurait au moins 6 symptômes d’inattention (9/9) et / ou au moins 6 symptômes  d’hyperactivité-impulsivité (9/9) présents depuis au moins 6 mois. ",
        "",
        ""
    )

    ajouter_paragraphe(
        doc,
        "",
        "Critère B : Il y aurait plusieurs symptômes (au moins 3) présents avant l’âge de 12 ans. Critère C et D : Les symptômes et la gêne sont observés dans au moins deux environnements de  fonctionnement.  ",
        ""
    )

    ajouter_paragraphe(
        doc,
        "",
        "Critère E : Les symptômes ne peuvent pas être (mieux) expliqués par la présence d’un autre trouble  psychiatrique.",
        ""
        )

    # Tableau : trouble déficitaire de l'attention
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False  # Désactive le redimensionnement automatique

    # Style de bordure pour chaque cellule
    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Données
    entetes = [f"{prenom}, au travers de cet entretien, ne semble pas présenter un trouble déficitaire de l’attention avec ou sans  hyperactivité."]

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()
   

def evaluation_tdah(critere_a_1, critere_a_2, nom, prenom):
    # EVALUATION MANIFESTATION TDA/H
    EMT = doc.add_paragraph()
    run = EMT.add_run("EVALUATION MANIFESTATION TDA/H ")
    run.bold = True
    run.font.size = Pt(15)
    EMT.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ajouter_paragraphe(
        doc,
        "",
        f"Au regard des difficultés rapportées lors de l’anamnèse, il nous est apparu nécessaire de réaliser un entretien  auprès de {prenom}. Il est important de rappeler qu’il s’agit d’un entretien clinique et que les résultats du bilan  puissent être quelque peu différents. En effet, il est possible que {prenom} soit dans de meilleures conditions lors  de l’évaluation : en dualité, sans bruit parasite ni de passage environnant.  ",
        ""
    )

    doc.add_paragraph()

    # Young Diva
    Young = doc.add_paragraph()
    run = Young.add_run("Young -DIVA 5.0 ")
    run.bold = True
    run.underline = True
    run2 = Young.add_run(": Au niveau des critères selon le DMS-5 ")

    ajouter_paragraphe(
        doc,
        "Critère A : Il y aurait au moins 6 symptômes d’inattention (9/9) et / ou au moins 6 symptômes  d’hyperactivité-impulsivité (9/9) présents depuis au moins 6 mois. ",
        "",
        ""
    )

    ajouter_paragraphe(
        doc,
        "",
        "Critère B : Il y aurait plusieurs symptômes (au moins 3) présents avant l’âge de 12 ans. Critère C et D : Les symptômes et la gêne sont observés dans au moins deux environnements de  fonctionnement.  ",
        ""
    )

    ajouter_paragraphe(
        doc,
        "",
        "Critère E : Les symptômes ne peuvent pas être (mieux) expliqués par la présence d’un autre trouble  psychiatrique.",
        ""
        )

    # Tableau : trouble déficitaire de l'attention
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False  # Désactive le redimensionnement automatique

    # Données
    entetes = [f"{prenom}, au travers de cet entretien, ne semble pas présenter un trouble déficitaire de l’attention avec ou sans  hyperactivité."]

    bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

    # Remplir l’en-tête
    for i, titre in enumerate(entetes):
        cell = table.cell(0, i)
        para = cell.paragraphs[0]
        run = para.add_run(titre)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

    doc.add_paragraph()

#************************************************************************************************
# BILAN NEUROPSYCHOLOGIQUE
#************************************************************************************************

def bilan_neuropsychologique(nom, prenom, date):
    pass


def font_reset():
    # Police désirée
    font_name = "Times New Roman"
    font_size = Pt(11)

    # Parcourir tous les paragraphes
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = font_size

    # Parcourir toutes les cellules des tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = font_name
                        run.font.size = font_size

def alignement_reset():
    # Justifier tous les paragraphes hors tableaux
    for paragraph in doc.paragraphs:
        if "BILAN PSYCHOMÉTRIQUE" in paragraph.text or "------------------------------------------------" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


