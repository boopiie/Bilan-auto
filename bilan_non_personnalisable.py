from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm

def para(phrase, nom_para):
    nom_para = doc.add_paragraph(phrase)

def run(phrase, nom_para, bold_value = False):
    nom_para.add_run(phrase).bold = bold_value

def ajouter_paragraphe(doc, texte_gras1, texte_normal1, texte_gras2, texte_normal3=None, texte_gras3=None, texte_normal4=None, texte_gras4=None, texte_normal5=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run1 = para.add_run(texte_gras1)
    run1.font.size = Pt(12)
    run1.bold = True

    run2 = para.add_run(texte_normal1)
    run2.font.size = Pt(12)
    run2.bold = False

    run3 = para.add_run(texte_gras2)
    run3.font.size = Pt(12)
    run3.bold = True
    
    if texte_normal3:
        run4 = para.add_run(texte_normal3)
        run4.font.size = Pt(12)
        run4.bold = False
    
    if texte_gras3:
        run4 = para.add_run(texte_gras3)
        run4.font.size = Pt(12)
        run4.bold = True

    if texte_normal4:
        run4 = para.add_run(texte_normal4)
        run4.font.size = Pt(12)
        run4.bold = False

    if texte_gras4:
        run4 = para.add_run(texte_gras4)
        run4.font.size = Pt(12)
        run4.bold = True
    
    if texte_normal5:
        run4 = para.add_run(texte_normal5)
        run4.font.size = Pt(12)
        run4.bold = False

def ajouter_paragraphe_italique(doc, texte_gras1, texte_normal1, texte_gras2=None, texte_normal3=None, texte_gras3=None, texte_normal4=None, texte_gras4=None, texte_normal5=None, texte_gras5=None, texte_normal6=None, texte_gras6=None, texte_normal7=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run1 = para.add_run(texte_gras1)
    run1.font.size = Pt(12)
    run1.bold = True
    run1.italic = True

    run2 = para.add_run(texte_normal1)
    run2.font.size = Pt(12)
    run2.bold = False
    run2.italic = True

    if texte_gras2:
        run3 = para.add_run(texte_gras2)
        run3.font.size = Pt(12)
        run3.bold = True
        run3.italic = True

    if texte_normal3:
        run4 = para.add_run(texte_normal3)
        run4.font.size = Pt(12)
        run4.bold = False
        run4.italic = True

    if texte_gras3:
        run5 = para.add_run(texte_gras3)
        run5.font.size = Pt(12)
        run5.bold = True
        run5.italic = True

    if texte_normal4:
        run6 = para.add_run(texte_normal4)
        run6.font.size = Pt(12)
        run6.bold = False
        run6.italic = True

    if texte_gras4:
        run7 = para.add_run(texte_gras4)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
    
    if texte_normal5:
        run1 = para.add_run(texte_normal5)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True
    
    if texte_gras5:
        run7 = para.add_run(texte_gras5)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
    
    if texte_normal6:
        run1 = para.add_run(texte_normal6)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True
    
    if texte_gras6:
        run7 = para.add_run(texte_gras6)
        run7.font.size = Pt(12)
        run7.bold = True
        run7.italic = True
    
    if texte_normal7:
        run1 = para.add_run(texte_normal7)
        run1.font.size = Pt(12)
        run1.bold = False
        run1.italic = True


def set_cell_border(cell, **kwargs):
    """
    Appliquer les bordures à une cellule de tableau.
    kwargs: top, bottom, start, end = {'sz': 12, 'val': 'single', 'color': '000000'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'bottom', 'start', 'end'):
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in kwargs[edge]:
                element.set(qn('w:{}'.format(key)), kwargs[edge][key])
            tcPr.append(element)
# Créer un nouveau document
doc = Document()

# En tête du texte
en_tete = doc.add_paragraph("Theo LOUSTAU")
en_tete.alignment = 1
en_tete = doc.add_paragraph("Psychologue - Neuropsychologue")
en_tete.alignment = 1
en_tete = doc.add_paragraph("Diplômé de l'Université de Bordeaux")
en_tete.alignment = 1
en_tete = doc.add_paragraph("283 rue Antoine Becquerel 40280 Saint Pierre-du-Mont")
en_tete.alignment = 1
en_tete = doc.add_paragraph("theoloustau.neurospy@gmail.com")
en_tete.alignment = 1
doc.add_paragraph('')
en_tete = doc.add_paragraph()
eval_neur = en_tete.add_run("ÉVALUATION NEUROPSYCHOLOGIQUE ")
eval_neur.bold = True
eval_neur.font.size = Pt(17)
en_tete.alignment = 1

# Premier cadre
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.cell(0, 0)
cadre1 = cell.add_paragraph()
texteBleu = cadre1.add_run("NOM et PRENOM                                                                "
                           "                    (NOM)  (Prenom) ")
texteBleu.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu.bold = True

cadre1 = cell.add_paragraph("Date de naissance                                                      "
"                                        (jj/mm/aaaa)")
cadre1 = cell.add_paragraph("Age au moment de l'évaluation                                                      "
"                 6 ans 4 mois ")
cadre1 = cell.add_paragraph("Latéralité                                                      "
"                                                             (lateralité)")
cadre1 = cell.add_paragraph("Date du bilan                                                       "
"                                                              (date)")
cadre1.add_run("\n")

doc.add_paragraph()

# Indication
indication = doc.add_paragraph()
IndicationBleu = indication.add_run("INDICATION, PLAINTE PRINCIPALE")
IndicationBleu.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
indication = doc.add_paragraph()
indication.add_run("J’ai rencontré ")
indication.add_run("NOM PRENOM ").bold = True
indication.add_run("à la demande de la PCO, afin de mieux comprendre son  fonctionnement cognitif.  ")
indication = doc.add_paragraph()
run("Ce présent bilan a donc pour objectif de définir le profil cognitif et comportemental de Léo, afin de fournir  des axes de " \
"travail et d'accompagnement. ", indication)

# Source information
source_info = doc.add_paragraph()
SourceBleu = source_info.add_run("SOURCE D'INFORMATION")
SourceBleu.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
source_info = doc.add_paragraph()   
run("Les informations pertinentes ont été recueillies auprès des parents et de Léo. ", source_info)
source_info = doc.add_paragraph()
source_info.add_run("Tests utilisés ").bold = True
run(": WISC-V ; TEA-Ch ; NEPSY-II ; KiTAP ; BROWN ; Young-DIVA ", source_info)

# Anamnese
anamnese = doc.add_paragraph()
anamnese.alignment = 1
titre = anamnese.add_run("ANAMNESE")
titre.bold = True
titre.font.size = Pt(14)
anamnese = doc.add_paragraph()

# Motif
anamnese.add_run("Motif de consultation :  ").bold = True
anamnese = doc.add_paragraph()
run("Léo est adressé par le PCO pour la réalisation d'un ", anamnese)
run("bilan neuropsychologique", anamnese, True)
run(", dans le cadre d'une", anamnese)
run("suspicion  de TDA/H", anamnese, True)
run(". Les difficultés sont observées aussi bien ", anamnese)
run("à la maison qu’à l’école", anamnese, True)
run(", où Léo présente un ", anamnese)
run("manque  de concentration", anamnese, True)
run(", une ", anamnese)
run("grande agitation", anamnese, True)
run(", et ", anamnese)
run("interrompt fréquemment les échanges", anamnese, True)
run(" en coupant la parole.  Ces comportements nuisent à sa scolarité et à ses relations sociales. ", anamnese)
anamnese = doc.add_paragraph()

# Suivis et bilan antérieurs
run("Suivis ou bilans antérieurs :  ", anamnese, True)
anamnese = doc.add_paragraph()
run("Léo va bénéficier d'un ", anamnese)
run("bilan et d'un suivi en psychomotricité", anamnese, True)
run(" et a bénéficié d'un ", anamnese)
run("suivi avec une thérapeute  du langage depuis deux ans", anamnese, True)
run(". Un ", anamnese)
run("bilan orthophonique est actuellement en cours", anamnese, True)
run(", et un ", anamnese)
run("accompagnement  AESH a été demandé ", anamnese, True)
run("pour soutenir sa scolarité.", anamnese)
anamnese = doc.add_paragraph()

# Famille
run("Famille : ", anamnese, True)
anamnese = doc.add_paragraph()
run("Léo est enfant unique et vit avec ses deux parents, qui sont ensemble. Son père est ", anamnese)
run("boucher/charcutier/traiteur", anamnese, True)
run(", et sa mère travaille dans le ", anamnese)
run("planning pour de l’aide à domicile. ", anamnese, True)
anamnese = doc.add_paragraph

# Bilan psychométrique

# Cadre
table2 = doc.add_table(rows=1, cols=1)

cell2 = table2.cell(0, 0)
cadre2 = cell2.add_paragraph()
cadre2.alignment = 1
texteGras = cadre2.add_run("Fonctionnement intellectuel global (WISC V)")
texteGras.bold = True
# Capacité cognitives globales
CapCo = doc.add_paragraph()
CapCo = doc.add_paragraph()
# Texte bleu
texteBleu2 = CapCo.add_run("Capacités cognitives globales ")
texteBleu2.bold = True
texteBleu2.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu2.underline = True
CapCo = doc.add_paragraph()
texteWISC = CapCo.add_run("Le WISC-V est utilisé pour mesurer les habiletés générales de raisonnement des " \
"enfants de 6 à 16 ans. Cette  évaluation fournit un score qui représente la capacité intellectuelle globale" \
" de l’enfant (QIT), ainsi que des  scores d’indice qui mesurent les domaines suivants du fonctionnement" \
" cognitif : compréhension verbale  (ICV), traitement visuospatial (IVS), raison fluide (IRF), mémoire de" \
" travail (IMT) et vitesse de traitement  (IVT). ")
texteWISC.italic = True
CapCo = doc.add_paragraph()
run("L’évaluation intellectuelle réalisée à l’aide du WISC-V met " \
"en évidence un profil présentant des capacités  intellectuelles hétérogènes." \
" En effet, l’hétérogénéité significative de son profil ne nous permet pas de " \
"calculer  un QIT chez Léo. En effet, des différences statistiquement significatives " \
"apparaissent entre plusieurs scores  d’indices. Alors, la note d’échelle totale"
" (QIT) – qui représente les aptitudes intellectuelles globales – ne peut " \
" nous permettre de comprendre le fonctionnement de Léo. L’étude des forces " \
"et des faiblesses est préconisée  pour mieux comprendre son profil cognitif. ", CapCo)

# Ajouter un titre centré et en italique
titre = doc.add_paragraph("Synthèse des notes composites principales")
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
titre.runs[0].italic = True

#
# Add table with 7 columns and 7 rows
table = doc.add_table(rows=7, cols=7)
table.style = 'Table Grid'
table.autofit = True

# Set column widths (approximate)
col_widths = [Inches(2.2), Inches(0.9), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.4), Inches(1.2)]
for row in table.rows:
    for idx, cell in enumerate(row.cells):
        cell.width = col_widths[idx]

# Header row
headers = [
    "Composite", "Indices", "Ensemble des Note Standard", "Note Composite",
    "Rang Percentile", "Intervalle de Confiance", "Description qualitative"
]
for i, text in enumerate(headers):
    p = table.cell(0, i).paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
data = [
    ["Compréhension Verbale", "ICV", "21", "103", "58", "94-112", "Moyenne"],
    ["Visuospatial", "IVS", "22", "105", "63", "97-112", "Moyenne"],
    ["Raisonnement Fluide", "IRF", "22", "106", "66", "98-113", "Moyenne"],
    ["Mémoire de travail", "IMT", "15", "85", "16", "78-95", "Moyenne faible"],
    ["Vitesse de traitement", "IVT", "15", "86", "18", "79-97", "Moyenne faible"],
    ["Échelle Totale", "QIT", "68", "98", "45", "92-104", "Moyenne"]
]

# Fill in data rows
for i, row_data in enumerate(data):
    row_idx = i + 1
    for j, val in enumerate(row_data):
        cell = table.cell(row_idx, j)
        p = cell.paragraphs[0]
        run = p.add_run(val)

        # Styling
        if row_idx == 6:
            # Last row background color
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "DCE6F1")
            cell._tc.get_or_add_tcPr().append(shading_elm)
        if row_idx in [4, 5]:
            run.font.color.rgb = RGBColor(255, 102, 0)
        if row_idx == 6 and j == 6:
            run.font.strike = True
# Description après tableau
ajouter_paragraphe(
    doc,
    "L'indice complémentaire d'aptitude générale (IAG = 104, RP = 61)",
    ", témoigne d'une capacité à raisonner,de facultés de compréhension, situées dans la ",
    "moyenne ",
    "de son âge."
)
doc.add_paragraph()
ajouter_paragraphe(
    doc,
    "L'indice de compétence cognitive (ICC = 82, RP = 12) ",
    "relatif aux traitements de bas niveaux (vitesse detraitement et mémoire de travail) se situe dans la ",
    "moyenne faible ",
    "de son âge."
)
doc.add_paragraph()
ajouter_paragraphe(
    doc,
    "L'indice non verbal (INV = 97, RP = 42) ",
    "relatif aux compétences non verbales se situe dans la ",
    "moyenne ",
    "de son âge."
)
doc.add_paragraph()
ajouter_paragraphe(
    doc,
    "",
    "Nous détaillerons, dans les chapitres suivants, les différents domaines cognitifs qui ont été évaluées," \
    " et qui permettent d'appréhender de manière plus approfondie le fonctionnement actuel de Léo.",
    "",
    ""
)
CapVe = doc.add_paragraph()
# Texte bleu
# Capacités verbales
texteBleu3 = CapVe.add_run("Capacités verbales")
texteBleu3.bold = True
texteBleu3.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu3.underline = True

ajouter_paragraphe_italique(
    doc,
    "L'indice de Compréhension Verbale (ICV) ",
    "mesure les aptitudes verbales en sollicitant le raisonnement, la compréhension, et la catégorisation. Il évalue la formation de concepts verbaux et les connaissances culturelles acquises dans l'environnement de l'enfant."
)

ajouter_paragraphe(
    doc,
    "",
    "Sur le plan qualitatif, nous observons que Léo comprend bien les consignes au décours du bilan. Il n'aura pas eu besoin d'énormément de reformulations ou d'explications supplémentaires. Le discours spontané est fluent, informatif et cohérent.",
    "",
    ""
)

ajouter_paragraphe(
    doc,
    "",
    "La note composite de compréhension verbale ",
    "(ICV = 103, RP = 58) ",
    "se situe dans la ",
    "moyenne ",
    "comparaison aux enfants du même âge."
)

# Tableau : Epreuve, Notes standards, Rang Percentile
# Création du tableau
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Données
entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
matrices = ["Similitudes", "10", "50"]
balances = ["Vocabulaire", "12", "75"]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (vert + italique)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 128, 0)  # vert
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

ajouter_paragraphe(
    doc,
    "",
    "Le subtest « ",
    "Similitudes ",
    "» permet d'appréhender la qualité du langage oral de Léo, ses capacités d'abstractionet de raisonnement verbal. Léo présente des performances dans la ",
    "moyenne forte ",
    "de son âge. Il montre de bonnes capacités catégorisation et de conceptualisation."
)

ajouter_paragraphe_italique(
    doc,
    "Exemple : « En quoi le cochon et le mouton se ressemblent ? Qu’est-ce qui fait qu’ils sont pareil ? »",
    ""
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de l'épreuve de «",
    "Vocabulaire",
    "», qui fait appel à ses connaissance internalisées (faisant appel à son expérience et des situations de la vie quotidienne), Léo obtient des résultats dans la ",
    "moyenne ",
    "de son âge. Léo témoigne d'un bon stock lexical."
)
ajouter_paragraphe_italique(
    doc,
    "Exemple : « Qu'est-ce qu'une fourchette ? »",
    ""
)
doc.add_paragraph()
# Fonctions visuo-spatiales
# Texte bleu
FVS = doc.add_paragraph()
texteBleu4 = FVS.add_run("Fonctions visuo-spatiales")
texteBleu4.bold = True
texteBleu4.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu4.underline = True

ajouter_paragraphe_italique(
    doc,
    "L'indice Visuo-Spatial (IVS) ",
    ", mesure la capacité à analyser les détails visuels et comprendre les relations  visuo-spatiales afin de construire des dessins géométriques à partir d'un modèle. Cette habileté requiert un  raisonnement visuo-spatial, l’intégration et la synthèse de relations « partie-tout », l'attention aux détails  visuels et l'intégration visuo-motrice. "
)
doc.add_paragraph()
ajouter_paragraphe(
    doc,
    "",
    "Dans ce domaine ",
    "(IVS = 105, RP = 63) ",
    "Léo possède des capacités visuo-constructives, d’analyse visuo spatiale, et de résolution de problème dans la ",
    "moyenne ",
    "de son âge."
)

# Tableau Epreuve cubes puzzle
# Création du tableau
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Données
entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
matrices = ["Cubes", "10", "50"]
balances = ["Puzzles visuels", "12", "75"]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (vert + italique)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 128, 0)  # vert
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Texte avec image cube
# Texte du paragraphe avec styles
doc.add_paragraph()
para = doc.add_paragraph()
run1 = para.add_run("Dans le subtest des « ")
run2 = para.add_run("Cubes")
run2.bold = True
run3 = para.add_run(" », où il est demandé à Léo de reproduire des patterns visuels à l’aide de cubes bicolores, Léo obtient des résultats dans la ")

run4 = para.add_run("moyenne")
run4.bold = True
run5 = para.add_run(" de son âge. Léo montre de bonnes capacités dans l’organisation spatiale des modèles. Malgré de bon résultats, Léo présente des difficultés de motricité fine dans cette épreuve pour manipuler les cubes.")

para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Ajout de l'image à droite

# Ajouter un tableau invisible pour texte + image à côté
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True

# Colonne 1 : le paragraphe
cell1 = table.cell(0, 0)
cell1.text = ""  # vide pour insérer plus tard
cell1_paragraph = cell1.paragraphs[0]
for run in para.runs:
    new_run = cell1_paragraph.add_run(run.text)
    new_run.bold = run.bold
cell1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Colonne 2 : l'image
cell2 = table.cell(0, 1)
cell2_paragraph = cell2.paragraphs[0]
run_image = cell2_paragraph.add_run()
run_image.add_picture("images/image_rien.jpg", width=Inches(1.5))

# Paragraphe descriptif "Puzzles visuels"
ajouter_paragraphe(
    doc,
    "",
    "Lors du subtest ",
    "« Puzzles visuels »",
    ", où il lui est demandé de choisir trois pièces de puzzle qui, ensemble, reconstruiraient le modèle visuel, Léo présente des performances , dans la ",
    "moyenne forte ",
    "de son âge. Léo montre de bonnes capacités visuo-spatiale. Toutefois, nous observons des difficultés pour soutenir l’effort dans cette épreuve ainsi qu’une certaine précipitation dans ses réponses."
)

# Insérer l’image à droite (adaptée à ton style)
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Titre "Raisonnement fluide"
RF = doc.add_paragraph()
texteBleu5 = RF.add_run("Raisonnement fluide")
texteBleu5.bold = True
texteBleu5.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu5.underline = True

# Paragraphe explicatif en italique
ajouter_paragraphe_italique(
    doc,
    texte_gras1 = "L’indice de raisonnement fluide (IRF), ",
    texte_normal1 = "permet de mesurer la capacité de Léo à détecter la relation conceptuelle sous-jacente entre des images et à utiliser le raisonnement pour identifier et appliquer des règles. L’identification et l’application des relations conceptuelles dans l’IRF exigent un raisonnement inductif et  quantitatif, l’intelligence à grande échelle, le traitement simultané et la pensée abstraite. "
)

# Texte général IRF
ajouter_paragraphe(
    doc,
    "",
    "Les capacités de raisonnement fluides",
    " (IRF = 106, RP = 66)",
    " telles que mesurées par le WISC-V, apparaissent ce jour, dans la ",
    "moyenne ",
    "pour son âge."
)

# Création du tableau
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Données
entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
matrices = ["Matrices", "10", "50"]
balances = ["Balances", "12", "75"]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (vert + italique)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 128, 0)  # vert
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Le subtest « ",
    "Matrices ",
    "» sollicite les compétences visuo-spatiales, le raisonnement visuel, les capacités  d’induction (inférer une logique à partir de l’observation), les capacités de  déduction (généralisation d’une logique et application sur de nouveaux  éléments) ainsi qu’une démarche catégorielle (abstraction de traits  communs et de différences par comparaisons). Ces derniers se situent dans  la ",
    "moyenne ",
    "de son âge. Léo montre de bonnes capacités en logique d’ordre  visuo-spatiale. Toutefois, il présente une grande impulsivité dans cette  épreuve, ne prenant pas toujours le temps de réflexion nécessaire avant de répondre aux items. "
)

para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

ajouter_paragraphe(
    doc,
    "",
    "Pour le subtest « ",
    "Balances",
    " », il s’agit d’une tâche de logique  inductive et déductive pour laquelle le concept quantitatif d’égalité  doit être acquis afin de permettre l’application des concepts de  correspondance, d’addition et/ou de multiplication. Dans cette  épreuve, Léo se situe dans la ",
    "moyenne forte ",
    "de son âge. Il montre un bon raisonnement logico-mathématique."
)

# IMAGE
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Memoire de travail
MDT = doc.add_paragraph()
texteBleu6 = MDT.add_run("Mémoire de travail")
texteBleu6.bold = True
texteBleu6.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu6.underline = True

ajouter_paragraphe_italique(
    doc,
    "La mémoire de travail (IMT) ",
    "est la capacité à manipuler de l’information que l’on maintient en mémoire à  court terme (ex. garder en tête les étapes d’un calcul mental, tout en effectuant la tâche de calcul). La mémoire à court terme représente un ensemble de processus qui permet la gestion d’un flux d’informations,  et leur stockage temporaire. Cette mémoire correspond à la quantité d’informations que l’on peut maintenir  active (ex. retenir un numéro de téléphone avant de le noter sur un papier).  "
)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Les résultats à cet indice (",
    "IMT = 85, RP = 16",
    ") témoigne d’une mémoire de travail dans la ",
    "moyenne faible ",
    "de son âge"
)

doc.add_paragraph()

# Tableau memoire des chiffres et des images
# Création du tableau
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Données
entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
matrices = ["Mémoire des Chiffres ", "10", "50"]
balances = ["Mémoire des Images ", "12", "75"]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(237, 125, 49)  # orange
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (orange)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(237, 125, 49)  # orange
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

ajouter_paragraphe_italique(
    doc,
    "",
    "Le versant auditivo-verbale met en jeu la boucle phonologique afin de maintenir et manipuler les informations  auditivo-verbales pour la réalisation d’une tâche. Le versant visuo-spatiale met en jeu le calepin visuo-spatial  afin de maintenir et manipuler des informations visuelles pour réaliser la tâche. "
)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Lors du subtest « ",
    "Mémoire des chiffres ",
    "», Léo présente des performances dans la ",
    "moyenne faible ",
    "de son âge.  Léo présente de légères difficulté en mémoire à court terme sur un support auditif. Il semblerait que les tâches  sérielles et de même nature soient coûteuses pour lui. De plus, il présente de nombreux décrochages  attentionnels avec une attention qui fluctue pendant l’épreuve. Léo est debout, il fait du bruit et montre des  difficultés dans la compréhension des consignes."
)

ajouter_paragraphe(
    doc,
    "",
    "Lors du subtest « ",
    "Mémoire des images ",
    "», Léo se situe dans la ",
    "moyenne faible ",
    "de son âge. Léo présente des difficultés en mémoire à court terme sur un support visuel. Lors de cette épreuve, il arrive vite à saturation et  ne semble plus disposer des ressources attentionnelles nécessaire pour continuer l’épreuve."
)

# IMAGE
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Vitesse de traitement
VDT = doc.add_paragraph()
texteBleu6 = VDT.add_run("Vitesse de traitement")
texteBleu6.bold = True
texteBleu6.font.color.rgb = RGBColor(0, 100, 255)  # Couleur
texteBleu6.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "La vitesse de traitement de l’information (IVT) réfère au rythme auquel un individu passe à travers différentes  étapes de cognition lors de l’exécution d’une tâche. Ces subtests font appel à la discrimination visuelle, à  l’organisation perceptive mais aussi aux habiletés graphiques. "
)

ajouter_paragraphe(
    doc,
    "",
    "La vitesse de traitement (",
    "IVT = 86, RP = 18",
    "), telle que mesurée par le WISC-V, apparaissent ce jour dans la ",
    "moyenne faible",
    "de son âge."
)

# Tableau code et symbole
# Création du tableau
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Données
entetes = ["Épreuves", "Notes Standards", "Rang Percentile"]
matrices = ["Code", "10", "50"]
balances = ["Symboles", "12", "75"]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(237, 125, 49)  # orange
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (orange)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(237, 125, 49)  # orange
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Lors du subtest « ",
    "Code ",
    "» Léo présente des performances dans la ",
    "moyenne faible ",
    "comparativement aux enfants de son âge. Il montre  une certaine lenteur dans sa vitesse de traitement. De plus, il présente  de grandes difficultés grapho-motrices dans cette épreuve et se  disperse à plusieurs reprises, ce qui l’a pénalisé dans sa performance. " 
)

# IMAGE
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

ajouter_paragraphe(
    doc,
    "",
    "Lors de l’épreuve « ",
    "Symboles",
    " », Léo se situe dans la ",
    "moyenne faible ",
    "de son âge. Léo présente une lenteur dans sa vitesse de traitement. Il présente également des  difficultés pour barrer les symboles et perd du temps en regardant le chronomètre à plusieurs reprises."
)

# IMAGE
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Conclusion
paraConclusion = doc.add_paragraph()
runCclusion = paraConclusion.add_run("CONCLUSION")
runCclusion.bold = True
runCclusion.font.size = Pt(15)
paraConclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "L’évaluation psychométrique de Léo met en évidence ",
    "de bonnes capacités intellectuelles",
    ", notamment dans  les domaines du raisonnement verbal, visuo-spatial et du raisonnement fluide, qui se situent dans la moyenne  de son âge. Ses compétences verbales sont satisfaisantes, son raisonnement fluide ainsi que ses capacités  d’analyse visuo-spatiale sont également bien développés. Ces résultats témoignent d’un potentiel cognitif  préservé, malgré un profil globalement hétérogène. "
)

ajouter_paragraphe(
    doc,
    "",
    "Toutefois, les conditions d’évaluation ont été marquées par ",
    "une grande agitation motrice et une attention  très fluctuante, avec de nombreux décrochages ",
    "tout au long des épreuves. Léo s’est montré ",
    "distractible  par l’environnement",
    ", avec de fréquentes digressions verbales, rendant certaines tâches difficiles à mener  jusqu’à leur terme. Lors des épreuves plus coûteuses sur le plan attentionnel, telles que celles impliquant la  mémoire de travail ou la vitesse de traitement, ",
    "une saturation rapide ",
    "a été observée, traduisant une difficulté  à maintenir l’effort cognitif " \
    "dans la durée. Ces éléments ont pu impacter négativement certaines de ses  " \
    "performances, en particulier dans les domaines sollicitant les ressources " \
    "attentionnelles soutenues et les capacités de concentration. "
)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "En somme, les résultats de ce bilan suggèrent un ",
    "fonctionnement cognitif globalement satisfaisant ",
    "mais  entravé par des ",
    "difficultés majeures dans la régulation attentionnelle."
)

# INTERPREATION – ECHELLES BROWN EF/A 
paraConclusion = doc.add_paragraph()
runCclusion = paraConclusion.add_run("INTERPREATION – ECHELLES BROWN EF/A")
runCclusion.bold = True
runCclusion.font.size = Pt(15)
paraConclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER

ajouter_paragraphe(
    doc,
    "",
    "Les échelles Brown EF/A fournissent une Note Totale et six Notes Cluster. La Note Totale est la plus  exhaustive, puisqu’elle comprend tous les items du test et propose une idée générale du fonctionnement  exécutif. Les six Notes Cluster rendent comptent d’aspects plus spécifiques des fonctions exécutives et offrent  un état des lieux plus ciblés des troubles du patient. ",
    ""
)

doc.add_paragraph()

ajouter_paragraphe_italique(
    doc,
    "Le Cluster 1. Activation",
    ", correspond à la capacité d’organiser, de prioriser et à démarrer le travail. ",
    "Le Cluster 2. Focus",
    ", correspond à la capacité à soutenir l’attention et à se concentrer sur des tâches. ",
    "Le Cluster 3. Effort",
    ", correspond à la capacité à maintenir l’énergie, l’effort et à ajuster la vitesse de  traitement. ",
    "Le Cluster 4. Émotion",
    ", correspond à la capacité à gérer la frustration et à moduler les émotions. ",
    " Le Cluster 5. Mémoire",
    ", correspond à la capacité à utiliser la mémoire de travail et à accéder aux  connaissances apprises.  ",
    "Le Cluster 6. Action",
    ", correspond à la capacité à surveiller et à autoréguler l’action. "
)

# Tableau : Cluster
# Création du tableau : 4 lignes, 1 colonne
table = doc.add_table(rows=4, cols=1)
table.style = 'Table Grid'
table.autofit = False  # Désactive le redimensionnement automatique

# Contenu des lignes
lignes = [
    ("70 et plus", " Nettement atypique (problème très significatif)"),
    ("60 – 69", " Modérément atypique (problème significatif)"),
    ("55-59", " Légèrement atypique (problème potentiellement important)"),
    ("54 et moins", " Typique (problème peu significatif)")
]

# Style de bordure pour chaque cellule
bordure = {'sz': '12', 'val': 'single', 'color': '000000'}

# Remplir l’en-tête
for i, phrase in enumerate(lignes):
    cell = table.cell(i, 0)
    para = cell.paragraphs[0]
    run = para.add_run(phrase)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

# Tableau : Activation, Focus, etc
table = doc.add_table(rows=3, cols=8)
table.style = 'Table Grid'
table.autofit = False  # Désactive le redimensionnement automatique

# Données
entetes = ["", "Activation", "Focus", "Effort", "Émotion","Mémoire","Action","Total"]
matrices = ["Parents","71","70","76","61","69","69","72"]
balances = ["Enseignant(e)","71","70","76","61","69","69","72"]

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 1 : Matrices
for i, val in enumerate(matrices):
    cell = table.cell(1, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

# Ligne 2 : Balances (orange)
for i, val in enumerate(balances):
    cell = table.cell(2, i)
    para = cell.paragraphs[0]
    run = para.add_run(val)
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Les ",
    "parents ",
    "de Léo rapportent des difficultés :  "
)

ajouter_paragraphe(
    doc,
    "- Modérément atypiques ",
    "dans les clusters : Émotion ; Mémoire ; Action ",
    ""
)

ajouter_paragraphe(
    doc,
    "", 
    "Au total, ces résultats mettent en avant des ",
    "difficultés exécutives nettement atypiques."
)

# IMAGE
para_img = doc.add_paragraph()
run_img = para_img.add_run()
run_img.add_picture("images/image_rien.jpg", width=Inches(1.5))  # Ajuste la taille si besoin
para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()


# EVALUATION MANIFESTATION TDA/H 
EMT = doc.add_paragraph()
run = EMT.add_run("EVALUATION MANIFESTATION TDA/H ")
run.bold = True
run.font.size = Pt(15)
EMT.alignment = WD_ALIGN_PARAGRAPH.CENTER

ajouter_paragraphe(
    doc,
    "",
    "Au regard des difficultés rapportées lors de l’anamnèse, il nous est apparu nécessaire de réaliser un entretien  auprès de Léo. Il est important de rappeler qu’il s’agit d’un entretien clinique et que les résultats du bilan  puissent être quelque peu différents. En effet, il est possible que Léo soit dans de meilleures conditions lors  de l’évaluation : en dualité, sans bruit parasite ni de passage environnant.  ",
    ""
)

doc.add_paragraph()

# Young Diva
Young = doc.add_paragraph()
run = Young.add_run("Young -DIVA 5.0 ")
run.bold = True
run.underline = True
run2 = Young.add_run(": Au niveau des critères selon le DMS-5 ")

ajouter_paragraphe(
    doc,
    "Critère A : Il y aurait au moins 6 symptômes d’inattention (9/9) et / ou au moins 6 symptômes  d’hyperactivité-impulsivité (9/9) présents depuis au moins 6 mois. ",
    "",
    ""
)

ajouter_paragraphe(
    doc,
    "",
    "Critère B : Il y aurait plusieurs symptômes (au moins 3) présents avant l’âge de 12 ans. Critère C et D : Les symptômes et la gêne sont observés dans au moins deux environnements de  fonctionnement.  ",
    ""
)

ajouter_paragraphe(
    doc,
    "",
    "Critère E : Les symptômes ne peuvent pas être (mieux) expliqués par la présence d’un autre trouble  psychiatrique.",
    ""
    )

# Tableau : trouble déficitaire de l'attention
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
table.autofit = False  # Désactive le redimensionnement automatique

# Données
entetes = ["Léo, au travers de cet entretien, ne semble pas présenter un trouble déficitaire de l’attention avec ou sans  hyperactivité."]

# Remplir l’en-tête
for i, titre in enumerate(entetes):
    cell = table.cell(0, i)
    para = cell.paragraphs[0]
    run = para.add_run(titre)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, top=bordure, bottom=bordure, start=bordure, end=bordure)

doc.add_paragraph()

# BILAN NEUROPSYCHOLOGIQUE 
BN = doc.add_paragraph()
run = BN.add_run("BILAN NEUROPSYCHOLOGIQUE ")
run.bold = True
run.font.size = Pt(15)
BN.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Capacités attentionelles
CA = doc.add_paragraph()
run = CA.add_run("Capacités attentionnelles : ")
run.bold = True
run.underline = True

ajouter_paragraphe(
    doc,
    "Alerte phasique / vigilance :  ",
    "",
    ""
)

ajouter_paragraphe_italique(
    doc,
    "",
    "L’alerte phasique concerne la capacité d’augmenter le niveau général d’attention en prévision d’un  événement attendu. ",
)

doc.add_paragraph()

# Tableau : Epreuve percentile sorciere kitap

table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'

# Ligne 1 – en-têtes
table.cell(0, 0).text = "Épreuve"
table.cell(0, 1).text = "Percentile"

epreuve = table.cell(0, 0)
epreuve.vertical_alignment

# Ligne 2 – Recherche dans le ciel
table.cell(1, 0).text = "La sorcière (KiTAP)"
p = table.cell(1, 1).paragraphs[0]
p.add_run("Temps : ")
r1 = p.add_run("Pc <1")
r1.bold = True
r1.font.color.rgb = RGBColor(255,0,0)
p.add_run("\nNote d’attention : ")
r3 = p.add_run("Pc <1")
r3.bold = True
r3.font.color.rgb = RGBColor(255,0,0)

doc.add_paragraph()

# La sorcière (KiTAP)
KiTAP = doc.add_paragraph()
run = KiTAP.add_run("La sorcière (KiTAP) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "l’épreuve évalue le temps de réaction (TR) lors d’une tâche simple, ce qui permet d'estimer la rapidité  générale de Léo et sa capacité à maintenir un niveau de réactivité stable (« alerte intrinsèque »). Ici, une  sorcière apparaît dans une fenêtre et il s’agit, pour Léo, de la faire disparaître au plus vite en actionnant la  touche réponse. "
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo présente un temps de réaction lent. De plus, l’écart-type de Léo atteste de grandes  difficultés pour maintenir un niveau de réactivité stable. Léo ne parvient donc pas à maintenir un bon niveau  d’alerte intrinsèque lors de cette épreuve.  ",
    ""
)

ajouter_paragraphe(
    doc,
    "Attention sélective visuelle :  ",
    "",
    ""
)

ajouter_paragraphe_italique(
    doc,
    "",
    "L’attention sélective visuelle concerne le fait de trier les informations visuelles disponibles pour traiter celles  qui sont pertinentes pour l’activité en cours, en inhibant une réponse aux autres stimuli présentés. "
)


# Tableau : Epreuve percentile attention sélective visuelle
# Création du tableau : 3 lignes, 2 colonnes
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'

# Ligne 1 – en-têtes
table.cell(0, 0).text = "Épreuve"
table.cell(0, 1).text = "Percentile"

# Ligne 2 – Recherche dans le ciel
table.cell(1, 0).text = "Recherche dans le ciel (TEA-Ch)"
p = table.cell(1, 1).paragraphs[0]
p.add_run("Cibles identifiées : ")
r1 = p.add_run("Pc 46")
r1.bold = True
p.add_run("\nTemps par cible : ")
r2 = p.add_run("Pc 80")
r2.bold = True
p.add_run("\nNote d’attention : ")
r3 = p.add_run("Pc 75")
r3.bold = True

# Ligne 3 – Carte géographique
table.cell(2, 0).text = "Carte géographique (TEA-Ch)"
p = table.cell(2, 1).paragraphs[0]
p.add_run("Cibles identifiées : ")
r = p.add_run("Pc 41")
r.bold = True


# TEA-Ch
TEACh = doc.add_paragraph()
run = TEACh.add_run("Recherche dans le ciel (TEA-Ch) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « Recherche dans le ciel » de la TEA-Ch évalue l’attention sélective visuelle de Léo. Elle consiste  à repérer des cibles spécifiques parmi un ensemble de distracteurs sur une feuille. Léo doit trier les  informations présentées, en identifiant les cibles tout en ignorant les éléments non pertinents. Cette tâche  sollicite des compétences d’exploration visuelle, de discrimination rapide, et d’organisation spatiale, tout en  nécessitant un contrôle moteur pour encercler les bonnes réponses. Le score final prend en compte le nombre  de cibles identifiées correctement, les erreurs éventuelles, et le temps moyen par cible. "
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo identifie 17 cibles sur 20 attestant de bonnes capacités pour repérer des stimuli  cibles parmi des distracteurs de manière visuelle. En moyenne, Léo identifie une cible en 6.7 secondes,  mettant en avant une bonne vitesse de traitement. Léo présente une stratégie exploratoire satisfaisante (exploration horizontale). Au total, Léo présente de bonnes capacités d’attention sélective visuelle dans cette  épreuve. ",
    ""
)

# TEA-Ch
TEACh = doc.add_paragraph()
run = TEACh.add_run("Carte géographique (TEA-Ch) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « Carte géographique » de la TEA-Ch évalue l’attention sélective visuelle de Léo. Léo doit repérer  et encercler des cibles spécifiques dispersées sur une carte parmi un ensemble de distracteurs visuels. Cette  tâche exige une exploration visuelle méthodique, une discrimination rapide entre les cibles et les distracteurs,  et une capacité à maintenir une vitesse de traitement satisfaisante sur une minute. Les résultats de cette  épreuve permettent d’apprécier la vitesse de traitement et l’attention sélective visuelle, tout en tenant compte  de la stratégie exploratoire utilisée par Léo. "
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo identifie 14 stimuli cibles. Léo montre de bonnes aptitudes pour identifier les  stimuli cibles en ignorant les stimuli distracteurs. Toutefois, Léo n’est pas intéressé par l’épreuve et décide  d’entourer ce qu’il a envie plutôt que de respecter les consignes. Il a donc eu besoin d’être recentré à plusieurs  reprises. ",
    ""
)

ajouter_paragraphe(
    doc,
    "Attention soutenue : ",
    "",
    ""
)

ajouter_paragraphe_italique(
    doc,
    "",
    "L’attention soutenue concerne l’action d’orienter intentionnellement son attention sur un stimulus et de la  maintenir dans le temps."
)


# Tableau : Epreuve percentile coup de fusil
# Création du tableau : 3 lignes, 2 colonnes
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'

# Ligne 1 – entêtes
table.cell(0, 0).text = "Épreuve"
table.cell(0, 1).text = "Percentile"

# Ligne 2 – Coups de fusil
table.cell(1, 0).text = "Coups de fusil (TEA-Ch)"
p = table.cell(1, 1).paragraphs[0]
p.add_run("Total : ")
r = p.add_run("Pc 7")
r.bold = True
r.font.color.rgb = RGBColor(255, 0, 0)  # rouge

# Ligne 3 – La danse des fantômes
table.cell(2, 0).text = "La danse des fantômes (KiTAP)"
p = table.cell(2, 1).paragraphs[0]
p.add_run("Fausses : ")
r1 = p.add_run("Pc 34")
r1.bold = True
p.add_run("\nOmises : ")
r2 = p.add_run("Pc 27")
r2.bold = True


# Coup de fusil
CDF = doc.add_paragraph()
run = CDF.add_run("Coups de fusil (TEA-Ch) :  ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « Coups de fusil » de la TEA-Ch évalue l’attention soutenue auditive. L’enfant écoute une série de  tirs de vaisseaux et doit compter mentalement les tirs entendus tout au long de l’épreuve. Cette tâche exige  une concentration prolongée ainsi qu’une vigilance constant. Les résultats permettent d’évaluer la capacité  de l’enfant à mobiliser et maintenir son attention sur une tâche monotone et exigeante sur une période  prolongée. "
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo se montre en difficulté pour mobiliser ses ressources attentionnelles un long  moment. En effet, il parvient à restituer une bonnes réponses sur dix attestant de difficultés en attention  soutenue auditive. Cette épreuve est particulièrement coûteuse pour Léo qui décide à plusieurs reprises de ne  pas compter car c’est trop compliqué pour lui. ",
    ""
)

# La danse des fantomes
LDDF = doc.add_paragraph()
run = LDDF.add_run("La danse des fantômes (KiTAP) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « La danse des fantômes » de la KiTAP évalue l’attention soutenue visuelle de l’enfant. Dans cette  tâche, des fantômes apparaissent successivement dans les fenêtres d’un château. L’enfant doit réagir en  appuyant sur un bouton lorsque deux fantômes de la même couleur se succèdent. Cette épreuve sollicite la  concentration sur une période prolongée. Les résultats prennent en compte le nombre d’erreurs et  d’omissions, permettant de mesurer la constance de l’attention et la capacité de l’enfant à rester concentré  sur une tâche monotone et exigeante. "
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo présente de légères fragilités pour maintenir son attention sur toute la durée de  l’épreuve en réalisant des erreurs et des omissions mais cela reste convenable. En revanche, cliniquement, Léo  est fatigué à la fin de cette épreuve et montrer des difficultés pour soutenir l’effort. ",
    ""
)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "Attention divisée :  ",
    "",
    ""
)

ajouter_paragraphe_italique(
    doc,
    "",
    "L’attention divisée est la capacité à répartir ses ressources attentionnelles entre deux tâches ou stimuli  simultanés, afin de les traiter en parallèle."
)

# Création d’un tableau 5 lignes x 2 colonnes
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'

# Largeurs de colonnes (optionnel, ajustable dans Word sinon)
table.columns[0].width = Pt(300)
table.columns[1].width = Pt(300)

# Fonction pour ajouter du texte stylisé
def add_text(cell, text, bold=False, color=None):
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'orange':
        run.font.color.rgb = RGBColor(255, 140, 0)

# Ligne 1
table.cell(0, 0).text = "Épreuve"
table.cell(0, 1).text = "Percentile"

# Ligne 2
table.cell(1, 0).text = "Faire deux choses à la fois (TEA-Ch)"
add_text(table.cell(1, 1), "Coût de la double tâche : ", bold=False)
add_text(table.cell(1, 1), "Pc 5", bold=True, color='red')

# Ligne 3
table.cell(2, 0).text = "Écouter deux choses à la fois (TEA-Ch)"
add_text(table.cell(2, 1), "Somme des cibles identifiées : ", bold=False)
add_text(table.cell(2, 1), "Pc 10", bold=True, color='orange')

# Ligne 4
table.cell(3, 0).text = "Les hiboux (KiTAP)"
p = table.cell(3, 1).paragraphs[0]
p.add_run("TR auditif : ")
r1 = p.add_run("Pc 21")
r1.font.color.rgb = RGBColor(255, 140, 0)
p.add_run("\nTR visuel : ")
r2 = p.add_run("Pc <1")
r2.font.color.rgb = RGBColor(255, 0, 0)
p.add_run("\nFausses : ")
r3 = p.add_run("Pc 54")
r3.bold = True

# Ligne 5 (fusion cellule de gauche)
table.cell(4, 0).merge(table.cell(4, 0))
add_text(table.cell(4, 1), "Omises : ", bold=False)
add_text(table.cell(4, 1), "Pc 7", bold=True, color='red')

# Faire deux choses à la fois (TEA-Ch) 
FDCALF = doc.add_paragraph()
run = FDCALF.add_run("Faire deux choses à la fois (TEA-Ch) :")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « Faire deux choses à la fois » de la TEA-Ch évalue l’attention divisée de l’enfant. Il doit réaliser  simultanément deux tâches : une tâche auditive et une tâche visuelle. Cette double sollicitation demande à  l’enfant de répartir efficacement son attention entre deux modalités (auditive et visuelle) tout en maintenant  un bon niveau de performance dans chaque tâche. Les résultats permettent d’évaluer la capacité de l’enfant  à traiter des informations simultanées sans négliger l’une ou l’autre des tâches. "
)

ajouter_paragraphe_italique(
    doc,
    "",
    "Lors de cette épreuve, Léo ne parvient pas à réaliser les deux tâches à la fois. En effet, Léo présente des  difficultés pour la tâche visuelle et pour la tâche auditive. Ces résultats mettent en avant des difficultés d’attention divisée sur deux modalités différentes. "
)

doc.add_paragraph()

# Écouter deux choses à la fois (TEA-Ch) 
EDCALF = doc.add_paragraph()
run = EDCALF.add_run("Écouter deux choses à la fois (TEA-Ch) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "L’épreuve « Écouter deux choses à la fois » de la TEA-Ch évalue l’attention divisée dans le domaine auditif.  Léo doit simultanément écouter et se concentrer sur deux types de stimuli auditifs. Cette tâche exige de répartir  les ressources attentionnelles entre deux flux d’informations auditives et de maintenir une performance  équilibrée dans les deux tâches. Les résultats permettent d’évaluer la capacité à gérer et traiter plusieurs  informations auditives en parallèle. "
)

ajouter_paragraphe_italique(
    doc,
    "",
    "Lors de cette épreuve, Léo présente des difficultés pour traiter simultanément ces deux tâches auditives. En  effet, Léo présente des difficultés sur la première tâche auditive (retenir le nom de l’animal dans l’histoire) et  sur la seconde tâche auditive (compter les tirs de vaisseaux). Ces résultats mettent en évidence des difficultés pour diviser son attention sur deux stimuli auditifs à la fois.  "
)

# Les hiboux 
LH = doc.add_paragraph()
run = LH.add_run("Les hiboux (KiTAP) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "Dans cette épreuve, l’attention divisée est explorée à l'aide d'une double tâche, l'une en modalité visuelle et  l’autre en modalité auditive, l'enfant étant sollicité d’actionner la touche réponse à l’apparition d’une cible  tantôt visuelle tantôt auditive. "
)

ajouter_paragraphe_italique(
    doc,
    "",
    "Lors de cette épreuve, Léo présente des difficultés pour traiter simultanément ces deux tâches. Il présente une  légère lenteur dans son temps de réponse moyen en modalité auditive et un fort ralentissement en modalité  visuelle. Il réalise également de nombreuses omissions mettant en avant des difficultés importantes pour  diviser son attention sur deux modalités différentes. "
)

doc.add_paragraph()

# Distractibilité

ajouter_paragraphe(
    doc,
    "Distractibilité",
    "",
    "",
)

ajouter_paragraphe_italique(
    doc,
    "",
    "La distractibilité est la tendance à détourner involontairement son attention d'une tâche principale en raison  de stimuli internes ou externes non pertinents."
)

# Tableau distractibilité
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'

# Ligne 1 – en-têtes
t1 = table.cell(0, 0).paragraphs[0]
t1.add_run("Épreuve")
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1.add_run("\n")
t1.add_run("\n")
t1 = table.cell(0, 1).paragraphs[0]
t1.add_run("Percentile")
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER

epreuve = table.cell(0, 0)
epreuve.vertical_alignment

# Ligne 2 – Recherche dans le ciel
fantome = table.cell(1, 0).paragraphs[0]
fantome.add_run("\n")
fantome.add_run("\n")
fantome.add_run("\n")
fantome.add_run("Le fantôme triste et le fantôme joyeux (KiTAP)")
p = table.cell(1, 1).paragraphs[0]

p.add_run("Avec distracteur :")
p.add_run("\nTemps : ")
r1 = p.add_run("Pc 24")
r1.bold = True
r1.font.color.rgb = RGBColor(35,25,25)
p.add_run("\n Fausses : ")
r2 = p.add_run("Pc 46")
r2.bold = True
p.add_run("\nOmises : ")
r3 = p.add_run("Pc 4")
r3.bold = True
r3.font.color.rgb = RGBColor(255,0,0)

p.add_run("\n")

p.add_run("\nSans distracteur : ")
p.add_run("\nTemps : ")
r4 = p.add_run("Pc 24")
r4.bold = True
r4.font.color.rgb = RGBColor(35,25,25)
p.add_run("\n Fausses : ")
r5 = p.add_run("Pc 46")
r5.bold = True
p.add_run("\nOmises : ")
r6 = p.add_run("Pc 4")
r6.bold = True
r6.font.color.rgb = RGBColor(255,0,0)



# Le fantôme triste et le fantôme joyeux (KiTAP)
LFTELFJ = doc.add_paragraph()
run = LFTELFJ.add_run("Le fantôme triste et le fantôme joyeux (KiTAP) : ")
run.italic = True
run.underline = True

ajouter_paragraphe(
    doc,
    "",
    "Cette épreuve évalue la capacité de l’enfant à ignorer des stimuli distracteurs tout en maintenant une  attention focalisée sur une tâche principale. ",
    ""
)

doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, Léo réalise de nombreuses omissions avec et sans distracteurs mettant en avant des  difficultés pour maintenir une attention focalisée pendant l’épreuve. De plus, il présente le double d’omissions  en présence de distracteurs, attestant de difficultés pour ignorer les éléments distracteurs de son  environnement.  ",
    "",
    ""
)

doc.add_paragraph()
doc.add_paragraph()

ajouter_paragraphe(
    doc,
    "Inhibition",
    "",
    ""
)

ajouter_paragraphe_italique(
    doc,
    "",
    "L’inhibition est une forme de contrôle qui nous permet de résister aux habitudes ou automatismes, aux tentations, distractions ou interférences."
)

doc.add_paragraph()

# Tableau : La chauve-souris
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'

# Ligne 1 – en-têtes
t1 = table.cell(0, 0).paragraphs[0]
t1.add_run("Épreuve")
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1.add_run("\n")
t1.add_run("\n")
t1 = table.cell(0, 1).paragraphs[0]
t1.add_run("Note standard / Percentile")
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER

epreuve = table.cell(0, 0)
epreuve.vertical_alignment

# Ligne 2 – Recherche dans le ciel
fantome = table.cell(1, 0).paragraphs[0]
fantome.add_run("\n")
fantome.add_run("\n")
fantome.add_run("La chauve-souris (KiTAP)")
p = table.cell(1, 1).paragraphs[0]

p.add_run("Temps : ")
r1 = p.add_run("Pc 24")
r1.bold = True

r1.font.color.rgb = RGBColor(35,25,25)
p.add_run("\n Fausses : ")
r2 = p.add_run("Pc 46")
r2.bold = True

p.add_run("\nOmises : ")
r3 = p.add_run("Pc 4")
r3.bold = True
r3.font.color.rgb = RGBColor(255,0,0)

p.add_run("\nIndice de la prestation totale : ")
r4 = p.add_run("Pc 7")
r4.bold = True
r4.font.color.rgb = RGBColor(255,0,0)

doc.add_paragraph()

# La chauve-souris (KiTAP)
LCS = doc.add_paragraph()
run = LCS.add_run("La chauve-souris (KiTAP) : ")
run.italic = True
run.underline = True

ajouter_paragraphe_italique(
    doc,
    "",
    "Cette épreuve évalue les capacités d’inhibition et de contrôle de l’impulsivité. Léo observe une série de stimuli apparaissant à l’écran (chauve-souris ou chat). Il doit réagir rapidement en appuyant sur le bouton uniquement lorsque la chauve-souris apparaît, et s’abstenir de répondre lorsque le chat est présenté."
)

ajouter_paragraphe(
    doc,
    "",
    "Lors de cette épreuve, malgré avoir répété à Léo qu’il fallait appuyer le plus vite possible sur le bouton réponse, il décide de prendre son temps pour ne pas se tromper. Malgré cela, il réalise quelques erreurs et une omission. Au vu du temps pris pour répondre, les résultats de Léo mettent tout de même en avant des difficultés pour contrôler un comportement automatique dans cette épreuve.",
    ""
)

doc.save("Bilan.docx") 


